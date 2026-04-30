"""Метрики по документу диссертации (страницы, источники, оформление)."""

from __future__ import annotations

import io
import re
from dataclasses import dataclass, field
from typing import Any, Iterator

from googleapiclient.http import MediaIoBaseDownload
from docx import Document  # type: ignore[import-untyped]
from docx.enum.text import WD_LINE_SPACING  # type: ignore[import-untyped]

from magister_checking.docs_extract import extract_plain_text, table_cell_content_blocks


# Порядок важен: более длинные фразы раньше (подстрочные совпадения в lower()).
# В т.ч. типичные заголовки: «СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ» (Гизатова),
# «СПИСОК ИСПОЛЬЗОВАННОЙ ЛИТЕРАТУРЫ» (ошибочный род — см. _bibliography_heading_issue_note),
# «ПАЙДАЛАНЫЛҒАН ӘДЕБИЕТТЕР» (қазақша).
# Без отдельного «литература» / «references» как подстрок: их поиск цепляется
# к «в литературе…» / «…references…»; отдельная строка — в _all_bibliography_section_starts.
_BIB_MARKERS = (
    "список использованной литературы",
    "список использованных источников",
    "пайдаланылған әдебиеттер",
    "список литературы",
    "использованная литература",
    "библиографический список",
)

_APPROVED_BIB_HEADING_LOWER = "список использованных источников"
_WRONG_GENITIVE_LITERATURE_HEADING_LOWER = "список использованной литературы"


def bibliography_heading_issue_note(plain: str) -> str | None:
    """Если в тексте есть неверный заголовок без утверждённого — предупреждение для Stage 4."""

    low = plain.lower()
    if _WRONG_GENITIVE_LITERATURE_HEADING_LOWER not in low:
        return None
    if _APPROVED_BIB_HEADING_LOWER in low:
        return None
    return (
        "Заголовок списка литературы: по методичке используйте «СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ» "
        "(в документе указано «СПИСОК ИСПОЛЬЗОВАННОЙ ЛИТЕРАТУРЫ»)."
    )

_CHARS_PER_PAGE_RU = 2200

# Номер пункта «1. / 1) / [1]» в библиографии; больше — id/фрагменты URL в plain.
# 1000 с запасом выше реальных магистерских списков; отсекает 1591, 27116051 и т.п.
_MAX_PLAUSIBLE_BIBLIO_INDEX = 1000


@dataclass
class DissertationMetrics:
    approx_pages: int
    pdf_pages: int | None
    sources_count: int | None
    review_pages: int | None
    review_sources_count: int | None
    has_literature_review: bool
    has_results: bool
    has_discussion: bool
    formatting_compliance: bool | None
    font_size_14_ratio: float | None
    times_new_roman_ratio: float | None
    single_spacing_ratio: float | None
    headings_found: list[str] = field(default_factory=list)
    notes: list[str] = field(default_factory=list)
    page_margins_cm: dict[str, float] | None = None
    """Поля «доминирующей» секции (top/bottom/left/right в см).

    Для DOCX — мода ``<w:pgMar>`` по всем ``<w:sectPr>`` в порядке
    конверсии twips ÷ 567. Для Google Doc — ``documentStyle.margin*`` (pt → см).
    ``None`` — поля не извлечены (нет ``sectPr`` / нестандартный API).
    """
    page_margins_secondary_cm: list[dict[str, float]] = field(default_factory=list)
    """Прочие комбинации полей, встретившиеся в документе (без основной).

    У Камзебаевой 7 секций имеют (1.83/2.12/1.75/0.75), 1 — (1.83/0.49/1.75/0.75)
    (титул), 1 — (2.5/2.12/1.75/0.75) (приложения). В ``page_margins_cm`` уйдёт
    мода, в ``secondary`` — оставшиеся две — для отчёта «у вас N секций
    с другими полями».
    """
    page_numbering_present: bool | None = None
    """True, если в документе обнаружено поле PAGE в footer-е, привязанном
    к большинству секций (или в основной body-final ``sectPr``).

    False — нумерации фактически нет (footer без PAGE, либо PAGE в footer-е,
    привязанном лишь к малой части секций — реальный кейс Камзебаевой:
    1 из 9 sectPr).
    None — не определено (например, у Google Doc API нет поля footers).
    """
    page_numbering_position: str | None = None
    """Положение номера страницы: 'bottom-left' / 'bottom-center' / 'bottom-right'
    / 'top-...' / None (если PAGE есть, но позиция не определена)."""
    page_numbering_sections_with_footer: int | None = None
    """Кол-во ``<w:sectPr>`` с явным footerReference (DOCX). Для Google Doc — None."""
    page_numbering_sections_total: int | None = None
    """Всего ``<w:sectPr>`` в документе (DOCX). Для Google Doc — None."""
    bibliography_heading_warning: str | None = None
    """Несовпадение с утверждённым названием раздела списка литературы (только текст предупреждения)."""


def iter_heading_texts(document: dict[str, Any]) -> Iterator[str]:
    """Тексты абзацев со стилем HEADING_* (включая ячейки таблиц)."""
    content = document.get("body", {}).get("content") or []
    if not isinstance(content, list):
        content = []
    yield from _headings_in_content(content)


def _headings_in_content(content: list[dict[str, Any]]) -> Iterator[str]:
    for element in content:
        if "paragraph" in element:
            p = element["paragraph"]
            st = p.get("paragraphStyle") or {}
            nst = st.get("namedStyleType") or ""
            if isinstance(nst, str) and nst.startswith("HEADING_"):
                t = _paragraph_text(p).strip()
                if t:
                    yield t
        elif "table" in element:
            for row in element["table"].get("tableRows", []):
                for cell in row.get("tableCells", []):
                    yield from _headings_in_content(table_cell_content_blocks(cell))


def _paragraph_text(paragraph: dict[str, Any]) -> str:
    parts: list[str] = []
    for pe in paragraph.get("elements", []):
        tr = pe.get("textRun")
        if tr and "content" in tr:
            parts.append(tr["content"])
    return "".join(parts)


def _iter_paragraphs_in_content(content: list[dict[str, Any]]) -> Iterator[dict[str, Any]]:
    for element in content:
        if "paragraph" in element:
            yield element["paragraph"]
        elif "table" in element:
            for row in element["table"].get("tableRows", []):
                for cell in row.get("tableCells", []):
                    yield from _iter_paragraphs_in_content(
                        table_cell_content_blocks(cell)
                    )


def iter_paragraphs(document: dict[str, Any]) -> Iterator[dict[str, Any]]:
    """Итерирует все paragraph-элементы документа, включая таблицы."""

    content = document.get("body", {}).get("content") or []
    if not isinstance(content, list):
        content = []
    yield from _iter_paragraphs_in_content(content)


def _count_styled_chars(text: str) -> int:
    return sum(1 for ch in text if not ch.isspace())


def _normalize_font_family(name: str | None) -> str:
    return re.sub(r"\s+", " ", str(name or "").strip().lower())


def _is_times_new_roman(name: str | None) -> bool:
    return _normalize_font_family(name) == "times new roman"


def _is_14_pt(magnitude: float | int | None, unit: str | None = None) -> bool:
    if magnitude is None:
        return False
    if unit and str(unit).upper() not in {"PT", ""}:
        return False
    return abs(float(magnitude) - 14.0) <= 0.1


def _is_single_spacing_percent(value: float | int | None) -> bool:
    if value is None:
        return True
    return abs(float(value) - 100.0) <= 0.5


def _safe_ratio(part: int, total: int) -> float | None:
    if total <= 0:
        return None
    return part / total


def _formatting_compliance(
    *,
    font_size_ratio: float | None,
    font_family_ratio: float | None,
    line_spacing_ratio: float | None,
) -> bool | None:
    ratios = (font_size_ratio, font_family_ratio, line_spacing_ratio)
    if any(r is None for r in ratios):
        return None
    return all((r or 0.0) > 0.95 for r in ratios)


def _google_named_styles(document: dict[str, Any]) -> dict[str, dict[str, Any]]:
    styles = (document.get("namedStyles") or {}).get("styles", [])
    out: dict[str, dict[str, Any]] = {}
    for style in styles:
        name = str(style.get("namedStyleType") or "").strip()
        if name:
            out[name] = style
    return out


def _google_effective_text_style(
    *,
    named_styles: dict[str, dict[str, Any]],
    paragraph: dict[str, Any],
    run_style: dict[str, Any] | None,
) -> dict[str, Any]:
    paragraph_style = paragraph.get("paragraphStyle") or {}
    named_style = str(paragraph_style.get("namedStyleType") or "NORMAL_TEXT")
    base = dict((named_styles.get(named_style) or {}).get("textStyle") or {})
    if run_style:
        base.update(run_style)
    return base


def _google_effective_paragraph_style(
    *,
    named_styles: dict[str, dict[str, Any]],
    paragraph: dict[str, Any],
) -> dict[str, Any]:
    paragraph_style = paragraph.get("paragraphStyle") or {}
    named_style = str(paragraph_style.get("namedStyleType") or "NORMAL_TEXT")
    base = dict((named_styles.get(named_style) or {}).get("paragraphStyle") or {})
    base.update(paragraph_style)
    return base


def _gdoc_dimension_to_cm(dim: dict[str, Any] | None) -> float | None:
    """Google Docs Dimension {'magnitude': float, 'unit': 'PT'} → см.

    API Google Docs всегда отдаёт margins в pt (см. v1 docs §Dimension).
    Если ``unit`` отсутствует или не PT — отдаём None, чтобы вызывающий
    мог корректно сказать «не определено», а не давать вводящие в
    заблуждение цифры.
    """

    if not dim:
        return None
    if dim.get("unit") not in (None, "PT"):
        return None
    return _pt_to_cm(dim.get("magnitude"))


def _gdoc_collect_section_margins(document: dict[str, Any]) -> list[dict[str, float]]:
    """Список ``{top,bottom,left,right}`` (см) по всем секциям Google Doc.

    Google Doc хранит margins в двух местах:
    - ``document.documentStyle.margin{Top,Bottom,Left,Right}`` — дефолт.
    - В ``content[].sectionBreak.sectionStyle.margin{Top,Bottom,Left,Right}``
      — переопределение для секции, начинающейся с этого ``SectionBreak``.

    По спецификации API: первая секция использует ``documentStyle``
    как дефолт; ``sectionStyle`` секции после ``SectionBreak``
    переопределяет дефолт. Возвращаем все встретившиеся комбинации
    (включая дефолт), порядок — как в документе.
    """

    out: list[dict[str, float]] = []
    keys = ("Top", "Bottom", "Left", "Right")

    def _collect(style: dict[str, Any] | None) -> dict[str, float] | None:
        if not style:
            return None
        margins: dict[str, float] = {}
        for k in keys:
            cm = _gdoc_dimension_to_cm(style.get(f"margin{k}"))
            if cm is not None:
                margins[k.lower()] = cm
        if {"top", "bottom", "left", "right"}.issubset(margins.keys()):
            return margins
        return None

    default = _collect(document.get("documentStyle"))
    if default is not None:
        out.append(default)

    for chunk in document.get("body", {}).get("content", []) or []:
        sb = chunk.get("sectionBreak") if isinstance(chunk, dict) else None
        if not sb:
            continue
        sec_style = sb.get("sectionStyle")
        m = _collect(sec_style)
        if m is not None:
            out.append(m)

    return out


def _gdoc_iter_footer_paragraphs(footer: dict[str, Any]) -> Iterator[dict[str, Any]]:
    for chunk in footer.get("content", []) or []:
        para = chunk.get("paragraph") if isinstance(chunk, dict) else None
        if isinstance(para, dict):
            yield para


def _gdoc_paragraph_alignment_to_position(alignment: str | None) -> str | None:
    """Google Doc ``paragraphStyle.alignment`` → 'left'/'center'/'right'.

    Возможные значения по API: ``START`` (left для LTR), ``CENTER``,
    ``END`` (right для LTR), ``JUSTIFIED``, ``ALIGNMENT_UNSPECIFIED``.
    Магистерские проекты у нас LTR (русский/казахский), так что
    START/END однозначно мапятся на left/right.
    """

    if not alignment:
        return None
    a = alignment.upper()
    if a in {"START", "ALIGNMENT_UNSPECIFIED"}:
        return "left"
    if a == "CENTER":
        return "center"
    if a == "END":
        return "right"
    if a == "JUSTIFIED":
        return "left"
    return None


def _gdoc_page_numbering_info(document: dict[str, Any]) -> dict[str, Any]:
    """Информация о нумерации страниц в Google Doc (через Docs API).

    Логика:
    - Перебираем ``document.footers`` (dict ``footerId -> Footer``).
    - В каждом footer ищем абзац, элементы которого содержат
      ``autoText.type == 'PAGE_NUMBER'``.
    - ``present = True`` если найден хотя бы один такой абзац.
    - ``position`` = ``bottom-{horizontal}`` по ``paragraphStyle.alignment``
      первого подходящего абзаца.

    У Google Doc нет аналога DOCX-секций с покрытием, поэтому
    ``sections_with_footer`` / ``sections_total`` всегда None — coverage-
    warning не генерируется (для GDoc это не имеет смысла, нумерация
    либо есть, либо нет, без частичности по разделам).
    """

    footers = document.get("footers") or {}
    if not isinstance(footers, dict) or not footers:
        return {
            "present": False,
            "position": None,
            "sections_with_footer": None,
            "sections_total": None,
        }

    for footer in footers.values():
        if not isinstance(footer, dict):
            continue
        for para in _gdoc_iter_footer_paragraphs(footer):
            elements = para.get("elements") or []
            has_page = False
            for el in elements:
                if not isinstance(el, dict):
                    continue
                auto = el.get("autoText") or {}
                if auto.get("type") == "PAGE_NUMBER":
                    has_page = True
                    break
            if not has_page:
                continue
            alignment = (para.get("paragraphStyle") or {}).get("alignment")
            horizontal = _gdoc_paragraph_alignment_to_position(alignment)
            position = f"bottom-{horizontal}" if horizontal else None
            return {
                "present": True,
                "position": position,
                "sections_with_footer": None,
                "sections_total": None,
            }

    return {
        "present": False,
        "position": None,
        "sections_with_footer": None,
        "sections_total": None,
    }


def _analyze_google_doc_formatting(document: dict[str, Any]) -> tuple[bool | None, float | None, float | None, float | None]:
    named_styles = _google_named_styles(document)
    total_chars = 0
    font_size_chars = 0
    font_family_chars = 0
    total_paragraphs = 0
    single_spacing_paragraphs = 0

    for paragraph in iter_paragraphs(document):
        text = _paragraph_text(paragraph)
        if text.strip():
            total_paragraphs += 1
            paragraph_style = _google_effective_paragraph_style(
                named_styles=named_styles,
                paragraph=paragraph,
            )
            if _is_single_spacing_percent(paragraph_style.get("lineSpacing")):
                single_spacing_paragraphs += 1

        for element in paragraph.get("elements", []):
            text_run = element.get("textRun") or {}
            content = str(text_run.get("content") or "")
            chars = _count_styled_chars(content)
            if chars <= 0:
                continue
            style = _google_effective_text_style(
                named_styles=named_styles,
                paragraph=paragraph,
                run_style=text_run.get("textStyle") or {},
            )
            total_chars += chars
            font_size = style.get("fontSize") or {}
            if _is_14_pt(font_size.get("magnitude"), font_size.get("unit")):
                font_size_chars += chars
            family = (style.get("weightedFontFamily") or {}).get("fontFamily") or style.get("fontFamily")
            if _is_times_new_roman(family):
                font_family_chars += chars

    font_size_ratio = _safe_ratio(font_size_chars, total_chars)
    font_family_ratio = _safe_ratio(font_family_chars, total_chars)
    line_spacing_ratio = _safe_ratio(single_spacing_paragraphs, total_paragraphs)
    return (
        _formatting_compliance(
            font_size_ratio=font_size_ratio,
            font_family_ratio=font_family_ratio,
            line_spacing_ratio=line_spacing_ratio,
        ),
        font_size_ratio,
        font_family_ratio,
        line_spacing_ratio,
    )


def _iter_docx_paragraphs(container: Any) -> Iterator[Any]:
    for paragraph in getattr(container, "paragraphs", []):
        yield paragraph
    for table in getattr(container, "tables", []):
        for row in table.rows:
            for cell in row.cells:
                yield from _iter_docx_paragraphs(cell)


def _docx_plain_text_all_paragraphs(doc: Any) -> str:
    """Плоский текст .docx: тело **и** ячейки таблиц (как в API Google Doc).

    Только ``Document.paragraphs`` не видит библиографию в таблицах — частый
    источник занижения/артефактов (max номера из обрывка списка).
    """
    parts: list[str] = []
    for p in _iter_docx_paragraphs(doc):
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    return "\n".join(parts)


_W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
_R_NS = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}"
_PKG_REL_NS = "{http://schemas.openxmlformats.org/package/2006/relationships}"

# 1 cm = 567 twips ровно (Word/OOXML — 1440 twips/inch, 2.54 cm/inch).
_TWIPS_PER_CM = 567.0
# 1 cm = 28.3464567 pt (72 pt/inch ÷ 2.54 cm/inch). Для Google Doc API
# (DocumentStyle.margin* — Dimension с ``unit='PT'``).
_PT_PER_CM = 28.3464567


def _twips_to_cm(value: str | None) -> float | None:
    if value is None:
        return None
    try:
        return round(int(value) / _TWIPS_PER_CM, 2)
    except (TypeError, ValueError):
        return None


def _pt_to_cm(magnitude: float | int | None) -> float | None:
    if magnitude is None:
        return None
    try:
        return round(float(magnitude) / _PT_PER_CM, 2)
    except (TypeError, ValueError):
        return None


def _docx_paragraph_numpr(paragraph: Any) -> tuple[str, str] | None:
    """Достаёт ``(numId, ilvl)`` из ``w:pPr/w:numPr`` абзаца, если задан.

    Это именно Word-нумерация (auto-list, «1.», «2.» отрисовываются Word'ом
    в рантайме). ``python-docx`` не возвращает её в ``paragraph.text``,
    поэтому без XML мы её не «видим». Возвращает ``None`` для абзацев без
    нумерованного списка.
    """

    elem = getattr(paragraph, "_p", None)
    if elem is None:
        return None
    npr = elem.find(f".//{_W_NS}numPr")
    if npr is None:
        return None
    nid = npr.find(f"{_W_NS}numId")
    ilvl = npr.find(f"{_W_NS}ilvl")
    nid_v = nid.get(f"{_W_NS}val") if nid is not None else None
    if nid_v is None:
        return None
    lvl_v = ilvl.get(f"{_W_NS}val") if ilvl is not None else "0"
    return (nid_v, lvl_v)


def _docx_paragraph_records(doc: Any) -> list[tuple[str, tuple[str, str] | None]]:
    """Линейный список ``(text, numpr)`` по всем абзацам тела и таблиц.

    Соответствует порядку склейки ``_docx_plain_text_all_paragraphs``,
    но включает абзацы с пустым текстом и сохраняет привязку к Word-нумерации.
    """

    out: list[tuple[str, tuple[str, str] | None]] = []
    for p in _iter_docx_paragraphs(doc):
        t = (p.text or "").strip()
        if not t:
            continue
        out.append((t, _docx_paragraph_numpr(p)))
    return out


def _is_bibliography_marker(text: str) -> bool:
    """True, если абзац — заголовок раздела «Список литературы» / «References» / …

    Унифицированная проверка для трёх независимых сигналов: word-list
    counter, line-numbering detector, URL-paragraph counter.
    """

    low = text.lower().strip(" :.\u2026")
    if low in {"литература", "references"}:
        return True
    return any(marker in low for marker in _BIB_MARKERS)


def _is_appendix_marker(text: str) -> bool:
    """True, если абзац — заголовок «Приложение …» / «ANNEX …» (граница хвоста)."""

    low = text.lower().lstrip()
    return low.startswith("приложение") or low.startswith("annex")


def _docx_bibliography_windows(
    doc: Any,
) -> list[list[tuple[str, tuple[str, str] | None]]]:
    """Все окна (text, numpr) от каждого маркера библиографии до «Приложение N».

    В одном документе может быть несколько вхождений маркера: например,
    у Танановой «СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ» встречается дважды —
    как ссылка в оглавлении (idx 36) и как реальный заголовок раздела
    (idx 318). Реальная библиография у неё **только** во втором окне,
    поэтому caller'ы (URL-paragraph counter, line-numbering detector)
    обходят все окна и берут максимум.

    Граница окна: первый абзац, начинающийся с «Приложение …» / «Annex …»
    (singular «Приложение»; plural «ПРИЛОЖЕНИЯ» намеренно НЕ ловится,
    т.к. это сводный заголовок раздела «приложений», который у Танановой
    стоит **между** TOC-вхождением маркера и реальной библиографией —
    обрыв в этой точке отсек бы реальный список).
    """

    records = _docx_paragraph_records(doc)
    if not records:
        return []
    bib_indices = [
        i for i, (text, _) in enumerate(records) if _is_bibliography_marker(text)
    ]
    if not bib_indices:
        return []
    windows: list[list[tuple[str, tuple[str, str] | None]]] = []
    for bib_idx in bib_indices:
        end = len(records)
        for j in range(bib_idx + 1, len(records)):
            if _is_appendix_marker(records[j][0]):
                end = j
                break
        windows.append(records[bib_idx + 1 : end])
    return windows


_DOCX_LINE_NUMBER_RE = re.compile(r"^\s*\d+[\.\)]\s*\S")
_DOCX_BRACKET_NUMBER_RE = re.compile(r"^\s*\[\d+\]\s*\S")


def _docx_bibliography_has_line_numbering(doc: Any) -> bool:
    """True, если хотя бы в одном окне библиографии ≥ 3 абзацев начинаются с «N.» / «[N]».

    Признак того, что text-эвристика max(n.) надёжна (Сулейменова/Мараджапова).
    Если ни одно окно не содержит реальной нумерации — text-индексы ловят
    шум (годы, диапазоны страниц), и тогда приоритет должен пойти к URL-counter
    или Word-списку.
    """

    for window in _docx_bibliography_windows(doc):
        matches = 0
        for text, _ in window:
            if _DOCX_LINE_NUMBER_RE.match(text) or _DOCX_BRACKET_NUMBER_RE.match(text):
                matches += 1
                if matches >= 3:
                    return True
    return False


_DOCX_URL_RE = re.compile(r"https?://", re.IGNORECASE)


def _docx_bibliography_url_paragraph_count(doc: Any) -> int | None:
    """Максимум URL-абзацев среди всех окон [маркер библиографии … «Приложение N»).

    Используется как (а) fallback при полном отсутствии нумерации и
    (б) корректировка Word-списка, если реальных записей с URL больше,
    чем длина auto-номерованного блока (см. ``analyze_docx_bytes`` —
    Тананова: word-list = 38, но фактически 43 записи, у каждой
    ровно один https://, остальные 5 — без auto-нумерации).

    Возвращает ``None`` для коротких библиографий (< 3 URL-абзацев),
    чтобы случайный URL в комментариях/предисловии не давал ложный сигнал.
    """

    windows = _docx_bibliography_windows(doc)
    if not windows:
        return None
    best = 0
    for window in windows:
        count = sum(1 for text, _ in window if _DOCX_URL_RE.search(text))
        if count > best:
            best = count
    return best if best >= 3 else None


def _docx_bibliography_word_list_count(doc: Any) -> int | None:
    """Считает длину Word-нумерованного списка, идущего сразу за маркером библиографии.

    Идея: после анкера «СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ» / «Литература» / …
    ищем максимальную **подряд** идущую группу абзацев с одинаковым
    ``(numId, ilvl)``; короткие разрывы (1 безномерный абзац) допускаются —
    бывает, что URL-абзац идёт без ``numPr``. Возвращает None, если такого
    списка нет (свободная вёрстка как у Танановой) — caller должен сделать
    fallback на text-индексы.

    Дополнительно считает «стрик» подряд идущих нумерованных абзацев с
    **любым** ``numPr`` (без non-numPr разрывов длиной > 1). Используется
    только **первый** такой стрик после маркера библиографии — он и есть
    реальный список источников. Последующие numPr-блоки (нумерованные
    пункты в приложениях/анкетах) игнорируются.

    Реальный кейс (Камзебаева, row 2): библиография разделена на kz-часть
    (numId=3, 42 пункта) и en-часть (numId=4, 19 пунктов), идущие впритык
    без пустых строк — визуально это один сквозной список из 61
    источника, но per-numPr максимум возвращал 42. Стрик-счётчик
    суммирует смежные numPr-блоки и даёт 61. После 61-го пункта идёт
    «Қосымша А» + текстовые строки анкеты (gap > 1) — стрик закрывается
    и нумерованные блоки анкеты (`('5','0')`, `('6','1')`, …) уже не
    учитываются. Для документов с одним numId стрик совпадает с
    per-numPr-максимумом (Гизатова, Досанов).
    """

    records = _docx_paragraph_records(doc)
    if not records:
        return None

    bib_para_indices: list[int] = []
    for idx, (text, _) in enumerate(records):
        if _is_bibliography_marker(text):
            bib_para_indices.append(idx)

    if not bib_para_indices:
        return None

    best: int | None = None
    for start_idx in bib_para_indices:
        run: dict[tuple[str, str], int] = {}
        active: tuple[str, str] | None = None
        active_count = 0
        gap = 0
        first_streak = 0
        streak_open = True
        for _, npr in records[start_idx + 1 :]:
            if npr is None:
                gap += 1
                if gap > 1:
                    if active is not None:
                        run[active] = max(run.get(active, 0), active_count)
                    active = None
                    active_count = 0
                    if first_streak > 0:
                        streak_open = False
                continue
            gap = 0
            if active is None or npr != active:
                if active is not None:
                    run[active] = max(run.get(active, 0), active_count)
                active = npr
                active_count = 1
            else:
                active_count += 1
            if streak_open:
                first_streak += 1
        if active is not None:
            run[active] = max(run.get(active, 0), active_count)

        if not run:
            continue
        candidate = max(max(run.values()), first_streak)
        # Слишком короткий «список» (1-2 пункта подряд) — это не библиография.
        if candidate < 3:
            continue
        if best is None or candidate > best:
            best = candidate

    return best


def _docx_collect_section_margins(doc: Any) -> list[dict[str, float]]:
    """Список ``{'top','bottom','left','right'}`` в см по всем ``<w:sectPr>``.

    Для каждой секции читает ``<w:pgMar>`` и конвертирует twips → см
    через ``_TWIPS_PER_CM`` (567 ровно). Секции без ``pgMar`` пропускает.
    Порядок результата соответствует порядку секций в документе (это
    нужно для последующего multi-section diagnose).
    """

    body_elem = getattr(getattr(doc, "element", None), "body", None)
    if body_elem is None:
        return []
    out: list[dict[str, float]] = []
    for sectpr in body_elem.iter(f"{_W_NS}sectPr"):
        pgmar = sectpr.find(f"{_W_NS}pgMar")
        if pgmar is None:
            continue
        margins: dict[str, float] = {}
        for key in ("top", "bottom", "left", "right"):
            cm = _twips_to_cm(pgmar.get(f"{_W_NS}{key}"))
            if cm is not None:
                margins[key] = cm
        if {"top", "bottom", "left", "right"}.issubset(margins.keys()):
            out.append(margins)
    return out


def _dominant_margins(
    margins_list: list[dict[str, float]],
) -> tuple[dict[str, float] | None, list[dict[str, float]]]:
    """Возвращает ``(мода, остальные уникальные комбинации)``.

    Если все секции одинаковые → ``(margins, [])``. Если есть несколько
    групп → берётся самая частая (мода); при ничьей — первая встретившаяся
    (детерминизм для тестов). Остальные уникальные комбинации (без моды)
    идут в ``secondary``, чтобы caller мог сообщить «у вас N секций
    с другими полями».
    """

    if not margins_list:
        return None, []
    keys = ("top", "bottom", "left", "right")
    counts: dict[tuple[float, ...], int] = {}
    order: list[tuple[float, ...]] = []
    for m in margins_list:
        key = tuple(m[k] for k in keys)
        if key not in counts:
            counts[key] = 0
            order.append(key)
        counts[key] += 1
    best_key = max(order, key=lambda k: counts[k])
    dominant = {k: v for k, v in zip(keys, best_key)}
    secondary = [
        {k: v for k, v in zip(keys, key)}
        for key in order
        if key != best_key
    ]
    return dominant, secondary


def _docx_footer_alignment_for_page_field(
    docx_bytes: bytes,
    footer_target: str,
) -> str | None:
    """Эффективное выравнивание абзаца с PAGE-полем в footer-файле.

    Читает ``word/<footer_target>``. Из всех ``<w:p>`` берёт первый,
    у которого есть ``<w:instrText>PAGE</w:instrText>`` или
    ``<w:fldSimple w:instr="PAGE"/>``. Возвращает ``<w:jc w:val="..."/>``
    напрямую из ``pPr`` (если задан) либо из применённого ``pStyle``
    (через ``word/styles.xml``). Если ничего не задано — возвращает
    ``'left'`` (дефолт OOXML).

    ``None`` — если footer-файла нет или PAGE-поля в нём нет.
    """

    import io as _io
    import zipfile

    try:
        with zipfile.ZipFile(_io.BytesIO(docx_bytes)) as zf:
            footer_path = f"word/{footer_target}"
            if footer_path not in zf.namelist():
                return None
            footer_xml = zf.read(footer_path)
            try:
                styles_xml = zf.read("word/styles.xml")
            except KeyError:
                styles_xml = None
    except (zipfile.BadZipFile, KeyError):
        return None

    from xml.etree import ElementTree as ET

    style_jc: dict[str, str] = {}
    if styles_xml is not None:
        try:
            sroot = ET.fromstring(styles_xml)
        except ET.ParseError:
            sroot = None
        if sroot is not None:
            for st in sroot.findall(f"{_W_NS}style"):
                sid = st.get(f"{_W_NS}styleId") or ""
                ppr = st.find(f"{_W_NS}pPr")
                if ppr is None:
                    continue
                jc_el = ppr.find(f"{_W_NS}jc")
                if jc_el is not None:
                    val = jc_el.get(f"{_W_NS}val")
                    if val:
                        style_jc[sid] = val

    try:
        froot = ET.fromstring(footer_xml)
    except ET.ParseError:
        return None
    for p in froot.iter(f"{_W_NS}p"):
        instr_texts = [
            (t.text or "").strip().upper()
            for t in p.iter(f"{_W_NS}instrText")
        ]
        simples = [
            (fs.get(f"{_W_NS}instr") or "").strip().upper()
            for fs in p.iter(f"{_W_NS}fldSimple")
        ]
        has_page = any("PAGE" in s for s in instr_texts) or any(
            "PAGE" in s for s in simples
        )
        if not has_page:
            continue
        ppr = p.find(f"{_W_NS}pPr")
        if ppr is not None:
            jc_el = ppr.find(f"{_W_NS}jc")
            if jc_el is not None:
                val = jc_el.get(f"{_W_NS}val")
                if val:
                    return val.lower()
            ps_el = ppr.find(f"{_W_NS}pStyle")
            if ps_el is not None:
                pstyle = ps_el.get(f"{_W_NS}val") or ""
                if pstyle in style_jc:
                    return style_jc[pstyle].lower()
        return "left"
    return None


def _jc_to_horizontal(jc: str | None) -> str | None:
    """OOXML ``<w:jc>`` → 'left' / 'center' / 'right' для отчёта.

    Принимаются исторические синонимы: ``start``=left, ``end``=right
    (LTR-документы; для RTL были бы наоборот, но магистерские проекты
    у нас на русском/казахском — оба LTR).
    """

    if not jc:
        return None
    j = jc.lower()
    if j in {"left", "start"}:
        return "left"
    if j in {"right", "end"}:
        return "right"
    if j in {"center", "centre"}:
        return "center"
    if j == "both":
        return "left"
    return None


def _docx_page_numbering_info(doc: Any, docx_bytes: bytes) -> dict[str, Any]:
    """Информация о нумерации страниц в DOCX.

    Алгоритм:
    1. Перебираем все ``<w:sectPr>``. Для каждой считаем
       ``<w:footerReference>`` (любого типа: default/first/even).
    2. Для секций с footerReference запоминаем ``r:id`` → разрешаем в
       ``word/_rels/document.xml.rels`` → ``footerN.xml``.
    3. Для каждого footer-файла проверяем наличие ``<w:instrText>PAGE</w:instrText>``
       (или ``<w:fldSimple w:instr="PAGE"/>``) и берём выравнивание.
    4. ``present = True`` если хотя бы один footer с PAGE привязан к более
       чем половине секций (или к единственной final-body секции).
       Иначе ``False`` — это случай Камзебаевой (1 sectPr с footer из 9).
    5. ``position`` — направление выравнивания PAGE-абзаца (``bottom-...``
       т.к. footer всегда внизу страницы; ``header_*`` отдельно не
       поддерживаем — у методички номер всегда внизу).

    Возвращает dict с ключами ``present``, ``position``,
    ``sections_with_footer``, ``sections_total``.
    """

    body_elem = getattr(getattr(doc, "element", None), "body", None)
    if body_elem is None:
        return {
            "present": None,
            "position": None,
            "sections_with_footer": None,
            "sections_total": None,
        }

    rels_map: dict[str, str] = {}
    import io as _io
    import zipfile
    from xml.etree import ElementTree as ET

    try:
        with zipfile.ZipFile(_io.BytesIO(docx_bytes)) as zf:
            try:
                rels_xml = zf.read("word/_rels/document.xml.rels")
            except KeyError:
                rels_xml = None
    except zipfile.BadZipFile:
        rels_xml = None
    if rels_xml is not None:
        try:
            rroot = ET.fromstring(rels_xml)
            for rel in rroot.findall(f"{_PKG_REL_NS}Relationship"):
                rid = rel.get("Id") or ""
                target = rel.get("Target") or ""
                rtype = rel.get("Type") or ""
                if "footer" in rtype.lower() and rid and target:
                    rels_map[rid] = target
        except ET.ParseError:
            pass

    sections_total = 0
    sections_with_footer = 0
    page_alignments: list[str] = []
    for sectpr in body_elem.iter(f"{_W_NS}sectPr"):
        sections_total += 1
        refs = sectpr.findall(f"{_W_NS}footerReference")
        if not refs:
            continue
        sections_with_footer += 1
        for ref in refs:
            rid = ref.get(f"{_R_NS}id")
            if not rid:
                continue
            target = rels_map.get(rid)
            if not target:
                continue
            align = _docx_footer_alignment_for_page_field(docx_bytes, target)
            if align is not None:
                page_alignments.append(align)

    if sections_total == 0:
        return {
            "present": None,
            "position": None,
            "sections_with_footer": 0,
            "sections_total": 0,
        }

    # Решение от 25.04.2026 (handoff §formatting_compliance): coverage сам по
    # себе — ненадёжный сигнал. У Мараджаповой 3/7 sectPr с явным
    # footerReference, и в Google Docs нумерация ВИДНА на каждой (наследует
    # default-footer). У Камзебаевой 1/9, и Google Docs нумерацию не
    # наследует — на большинстве страниц её визуально нет. По XML отличить
    # эти два случая надёжно нельзя. Поэтому ``present`` = «PAGE-поле есть
    # хотя бы в одном footer-файле, и хотя бы один sectPr явно указывает
    # footerReference». Coverage идёт в отчёт как warning, не блокирует.
    has_any_page = bool(page_alignments)
    present = bool(has_any_page and sections_with_footer > 0)

    position: str | None = None
    if page_alignments:
        # Берём первое успешное выравнивание (детерминизм для тестов;
        # на одном документе обычно все footer-PAGE одинаково выровнены).
        horizontal = _jc_to_horizontal(page_alignments[0])
        if horizontal is not None:
            position = f"bottom-{horizontal}"

    return {
        "present": present,
        "position": position,
        "sections_with_footer": sections_with_footer,
        "sections_total": sections_total,
    }


def _docx_font_size_pt(run: Any, paragraph: Any, document: Any) -> float | None:
    normal_style = document.styles["Normal"] if "Normal" in document.styles else None
    candidates = [
        getattr(getattr(run, "font", None), "size", None),
        getattr(getattr(getattr(run, "style", None), "font", None), "size", None),
        getattr(getattr(getattr(paragraph, "style", None), "font", None), "size", None),
        getattr(getattr(normal_style, "font", None), "size", None),
    ]
    for value in candidates:
        if value is not None:
            try:
                return float(value.pt)
            except Exception:  # noqa: BLE001
                continue
    return None


def _docx_font_name(run: Any, paragraph: Any, document: Any) -> str | None:
    normal_style = document.styles["Normal"] if "Normal" in document.styles else None
    candidates = [
        getattr(getattr(run, "font", None), "name", None),
        getattr(getattr(getattr(run, "style", None), "font", None), "name", None),
        getattr(getattr(getattr(paragraph, "style", None), "font", None), "name", None),
        getattr(getattr(normal_style, "font", None), "name", None),
    ]
    for value in candidates:
        if value:
            return str(value)
    return None


def _docx_is_single_spacing(paragraph: Any, document: Any) -> bool:
    normal_style = document.styles["Normal"] if "Normal" in document.styles else None
    fmt_candidates = [
        getattr(paragraph, "paragraph_format", None),
        getattr(getattr(paragraph, "style", None), "paragraph_format", None),
        getattr(normal_style, "paragraph_format", None),
    ]
    for fmt in fmt_candidates:
        if fmt is None:
            continue
        rule = getattr(fmt, "line_spacing_rule", None)
        if rule == WD_LINE_SPACING.SINGLE:
            return True
        value = getattr(fmt, "line_spacing", None)
        if value is None:
            continue
        if isinstance(value, (int, float)):
            return abs(float(value) - 1.0) <= 0.05
        return False
    return True


def _analyze_docx_formatting(document: Any) -> tuple[bool | None, float | None, float | None, float | None]:
    total_chars = 0
    font_size_chars = 0
    font_family_chars = 0
    total_paragraphs = 0
    single_spacing_paragraphs = 0

    for paragraph in _iter_docx_paragraphs(document):
        if (paragraph.text or "").strip():
            total_paragraphs += 1
            if _docx_is_single_spacing(paragraph, document):
                single_spacing_paragraphs += 1

        for run in getattr(paragraph, "runs", []):
            text = str(getattr(run, "text", "") or "")
            chars = _count_styled_chars(text)
            if chars <= 0:
                continue
            total_chars += chars
            if _is_14_pt(_docx_font_size_pt(run, paragraph, document)):
                font_size_chars += chars
            if _is_times_new_roman(_docx_font_name(run, paragraph, document)):
                font_family_chars += chars

    font_size_ratio = _safe_ratio(font_size_chars, total_chars)
    font_family_ratio = _safe_ratio(font_family_chars, total_chars)
    line_spacing_ratio = _safe_ratio(single_spacing_paragraphs, total_paragraphs)
    return (
        _formatting_compliance(
            font_size_ratio=font_size_ratio,
            font_family_ratio=font_family_ratio,
            line_spacing_ratio=line_spacing_ratio,
        ),
        font_size_ratio,
        font_family_ratio,
        line_spacing_ratio,
    )


def analyze_dissertation(document: dict[str, Any]) -> DissertationMetrics:
    plain = extract_plain_text(document)
    headings = list(iter_heading_texts(document))
    low = [h.lower() for h in headings]

    has_review = any(
        "обзор литературы" in h or "литературный обзор" in h for h in low
    )
    has_results = any("результат" in h for h in low)
    has_discussion = any(
        "обсуждение" in h or "вывод" in h for h in low
    )

    sources = _estimate_sources_count(plain)
    review_pages, review_sources = _estimate_review_metrics(plain)
    pages = max(1, len(plain) // _CHARS_PER_PAGE_RU)
    formatting_compliance, font_size_ratio, font_family_ratio, line_spacing_ratio = (
        _analyze_google_doc_formatting(document)
    )

    notes: list[str] = []
    if sources is None:
        notes.append(
            "Число источников не оценено: нет ожидаемого заголовка библиографии и/или "
            "нумерованного списка 1. … / 1) … / [1] …"
        )
    if formatting_compliance is None:
        notes.append("Соответствие оформлению не оценено (не хватило данных по стилям).")

    bib_heading_warn = bibliography_heading_issue_note(plain)
    if bib_heading_warn:
        notes.append(bib_heading_warn)

    margins_list = _gdoc_collect_section_margins(document)
    page_margins, page_margins_secondary = _dominant_margins(margins_list)
    numbering_info = _gdoc_page_numbering_info(document)

    return DissertationMetrics(
        approx_pages=pages,
        pdf_pages=None,
        sources_count=sources,
        review_pages=review_pages,
        review_sources_count=review_sources,
        has_literature_review=has_review,
        has_results=has_results,
        has_discussion=has_discussion,
        formatting_compliance=formatting_compliance,
        font_size_14_ratio=font_size_ratio,
        times_new_roman_ratio=font_family_ratio,
        single_spacing_ratio=line_spacing_ratio,
        headings_found=headings[:30],
        notes=notes,
        page_margins_cm=page_margins,
        page_margins_secondary_cm=page_margins_secondary,
        page_numbering_present=numbering_info.get("present"),
        page_numbering_position=numbering_info.get("position"),
        page_numbering_sections_with_footer=numbering_info.get("sections_with_footer"),
        page_numbering_sections_total=numbering_info.get("sections_total"),
        bibliography_heading_warning=bib_heading_warn,
    )


def _clip_tail_before_bibliography_postface(tail: str) -> str:
    """Обрезает хвост после нумерованного списка, если дальше идёт сводка по долям источников.

    Типично после списка: «Зарубежных источников 48% и отечественных источников 52%».
    Без обрезки в хвост попадает текст приложений/других нумерованных блоков, и
    max(n.) может раздуваться (реальный кейс: 42 ожидаемых vs 85 по шуму).
    """

    if len(tail) < 40:
        return tail
    # Доля с процентами в одной фразе с «источник» — раньше типичного «приложен…».
    patterns = (
        r"(?i)зарубежн(ых|ой|ие)\s+источник\w*\s+\d{1,3}\s*%",
        r"(?i)отечественн(ых|ой|ие)\s+источник\w*\s+\d{1,3}\s*%",
    )
    cut: int | None = None
    for pat in patterns:
        m = re.search(pat, tail)
        if m:
            pos = m.start()
            if cut is None or pos < cut:
                cut = pos
    if cut is not None and cut > 0:
        return tail[:cut]
    return tail


def _clip_tail_before_appendices(tail: str) -> str:
    """Обрезает хвост после библиографии, чтобы «Приложения…» / annex не портил счёт.

    Не используем подстроку ``приложен``: она входит в «прилож**ения**» (фраза
    «…источники и приложения») и **ложно** обрезала список в начале (реальный
    кейс: Гизатова — max оставался 29 вместо 106).
    """

    if len(tail) < 32:
        return tail
    low = tail.lower()
    cut = None
    for key in (
        "приложение",  # «ПРИЛОЖЕНИЕ 1», не «приложения» в «и приложения»
        "annex",
        "annexes",
    ):
        p = low.find(key, 16)
        if p != -1 and (cut is None or p < cut):
            cut = p
    if cut is not None and cut > 0:
        return tail[:cut]
    return tail


def _plausible_citation_index(n: int) -> bool:
    return 1 <= n <= _MAX_PLAUSIBLE_BIBLIO_INDEX


def _citation_index_numbers_in_text(tail: str) -> list[int]:
    """Номера, стоящие в начале строки: ``1. …``, ``1) …``, ``[1] …`` (многострочный текст)."""

    nums: list[int] = []
    for m in re.finditer(r"(?m)^\s*(\d+)[\.\)]\s*\S", tail):
        n = int(m.group(1))
        if _plausible_citation_index(n):
            nums.append(n)
    for m in re.finditer(r"(?m)^\s*\[(\d+)\]\s*\S", tail):
        n = int(m.group(1))
        if _plausible_citation_index(n):
            nums.append(n)
    return nums


_PAGE_RANGE_TAIL_RE = re.compile(r"\d+\s*[\-\u2013\u2014]\s*$")
_VOL_PAGE_TAIL_RE = re.compile(r"\d+\s*[:;]\s*$")


def _is_page_range_or_vol_page(tail: str, end_pos: int) -> bool:
    """Проверяет, что число — окончание диапазона ``453-459`` / тома ``15:194``.

    Анализирует короткий фрагмент непосредственно ПЕРЕД совпадением: если там
    «цифры + дефис/тире» (диапазон страниц) или «цифры + двоеточие» (том:страница) —
    это **не** индекс пункта библиографии (реальный кейс: Тананова, ``pp.453-459.``).
    """

    pre = tail[max(0, end_pos - 12) : end_pos]
    return bool(_PAGE_RANGE_TAIL_RE.search(pre) or _VOL_PAGE_TAIL_RE.search(pre))


def _citation_index_numbers_glued(tail: str) -> list[int]:
    """То же, но если Google API склеил абзацы без ``\\n`` (несколько ``n.`` в одной строке).

    Плюc подряд идущие вхождения (без привязки к началу всего хвоста). Годы
    1900–2100 отбрасываем, чтобы снизить шум в колонтитулах/датах. Также
    отбрасываем хвосты диапазонов страниц (``453-459.``) и тома (``15:194.``).
    """

    nums: list[int] = []
    for m in re.finditer(r"(\d+)[\.\)]\s+\S", tail):
        n = int(m.group(1))
        if 1900 <= n <= 2100:
            continue
        if not _plausible_citation_index(n):
            continue
        if _is_page_range_or_vol_page(tail, m.start()):
            continue
        nums.append(n)
    for m in re.finditer(r"\[(\d+)\]\s*\S", tail):
        n = int(m.group(1))
        if 1900 <= n <= 2100:
            continue
        if not _plausible_citation_index(n):
            continue
        nums.append(n)
    return nums


def _drop_upper_singleton_spike(sorted_unique: list[int]) -> list[int]:
    """Снимает верхний «остров» при скачке (…52, 53, 459): хвост приложения/шум.

    Порог: ``b - a > max(25, 2/3 * a)`` и ``a >= 3`` — не режем плотный ряд 98,99,100.
    """

    s = list(sorted_unique)
    while len(s) >= 2:
        a, b = s[-2], s[-1]
        gap = b - a
        if a >= 3 and gap > max(25, (a * 2) // 3):
            s.pop()
            continue
        break
    return s


def _tame_outlier_citation_max(nums: list[int]) -> int:
    """Если max — выброс относительно «основного» хвоста (1…63 и 485), берём второй max.

    Условие ``sechi >= 5``: иначе короткие списки (1, 2, 25) не ломаются.
    """

    s = sorted({n for n in nums if _plausible_citation_index(n)})
    if not s:
        return 0
    s = _drop_upper_singleton_spike(s)
    if len(s) < 2:
        return s[-1]
    hi, sechi = s[-1], s[-2]
    if sechi >= 5 and hi > 5 * sechi:
        return sechi
    return hi


def _max_citation_index_in_text_chunk(text: str) -> int | None:
    """Макс. индекс; сначала строки, начинающиеся с ``n.``; иначе glued. Выбросы режем."""

    nums = _citation_index_numbers_in_text(text)
    if nums:
        m = _tame_outlier_citation_max(nums)
        return m or None
    nums = _citation_index_numbers_glued(text)
    if not nums:
        return None
    m = _tame_outlier_citation_max(nums)
    return m or None


def _all_bibliography_section_starts(plain: str) -> list[int]:
    """Все подходящие начала «блока библиографии» (маркер или отдельная строка-заголовок)."""

    lower = plain.lower()
    starts: set[int] = set()
    for m in _BIB_MARKERS:
        start = 0
        while True:
            pos = lower.find(m, start)
            if pos < 0:
                break
            starts.add(pos)
            start = pos + max(1, len(m) // 2)
    for m in re.finditer(r"(?im)^\s*литература\s*$", plain):
        starts.add(m.start())
    for m in re.finditer(r"(?im)^\s*references\s*$", plain):
        starts.add(m.start())
    return sorted(starts)


def _max_sources_from_bib_start(plain: str, start: int) -> int | None:
    tail = plain[start:]
    tail = _clip_tail_before_bibliography_postface(tail)
    tail = _clip_tail_before_appendices(tail)
    return _max_citation_index_in_text_chunk(tail)


def _estimate_sources_count(plain: str) -> int | None:
    """Секция библиографии: max номер нумерованного пункта по **всем** подходящим якорям.

    Ранее брался только **самый правый** маркер; краткий блок «литература» в конце
    затирал длинный список. Берём max по всем вхождениям ``_BIB_MARKERS`` и
    отдельных строк *Литература* / *References*.

    Логика «1 … N → N» — как при ручной сверке по последней записи.
    """
    starts = _all_bibliography_section_starts(plain)
    if not starts:
        return None
    best: int | None = None
    for st in starts:
        n = _max_sources_from_bib_start(plain, st)
        if n is not None and (best is None or n > best):
            best = n
    return best


def count_pdf_pages_via_drive_export(*, drive_service: Any, file_id: str) -> int | None:
    """
    Экспортирует Google Doc в PDF через Drive API и считает страницы по маркерам в PDF.
    Без внешних зависимостей.
    """
    try:
        req = drive_service.files().export(fileId=file_id, mimeType="application/pdf")
        fh = io.BytesIO()
        downloader = MediaIoBaseDownload(fh, req)
        done = False
        while not done:
            _status, done = downloader.next_chunk()
        pdf = fh.getvalue()
        if not pdf:
            return None
        # Простейшая эвристика: /Type /Page встречается на каждую страницу,
        # а /Type /Pages — на корневой объект, его исключаем.
        pages = pdf.count(b"/Type /Page")
        pages -= pdf.count(b"/Type /Pages")
        return pages if pages > 0 else None
    except Exception:  # noqa: BLE001
        return None


def download_drive_file_bytes(*, drive_service: Any, file_id: str) -> bytes | None:
    """Скачивает файл с Google Drive (alt=media) в память."""
    try:
        req = drive_service.files().get_media(fileId=file_id)
        fh = io.BytesIO()
        downloader = MediaIoBaseDownload(fh, req)
        done = False
        while not done:
            _status, done = downloader.next_chunk()
        data = fh.getvalue()
        return data if data else None
    except Exception:  # noqa: BLE001
        return None


def _docx_page_count(docx_bytes: bytes) -> int | None:
    """
    Страницы из docProps/app.xml (Word считает страницы при сохранении).
    Это не «рендер» как PDF, но обычно ближе к реальности, чем оценка по символам.
    """
    try:
        import zipfile
        from xml.etree import ElementTree as ET

        with zipfile.ZipFile(io.BytesIO(docx_bytes)) as zf:
            xml = zf.read("docProps/app.xml")
        root = ET.fromstring(xml)
        # <Pages> может быть в пространстве имён, поэтому ищем по локальному имени
        for el in root.iter():
            if el.tag.endswith("Pages") and (el.text or "").strip().isdigit():
                return int((el.text or "0").strip())
        return None
    except Exception:  # noqa: BLE001
        return None


def analyze_docx_bytes(docx_bytes: bytes) -> DissertationMetrics:
    """
    Анализ Word (.docx):
    - pages: из docProps/app.xml (если есть), иначе оценка по символам
    - review_pages/review_sources: по секции «Обзор литературы» (по heading/text)
    """
    doc = Document(io.BytesIO(docx_bytes))
    plain = _docx_plain_text_all_paragraphs(doc)
    word_list_count = _docx_bibliography_word_list_count(doc)

    # headings: по стилю Word (Heading 1/2/… или русское "Заголовок")
    headings: list[str] = []
    para_info: list[tuple[str, bool]] = []
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if not text:
            continue
        style_name = ""
        try:
            style_name = str(p.style.name or "")
        except Exception:  # noqa: BLE001
            style_name = ""
        is_heading = bool(re.search(r"(?i)\bheading\b", style_name) or "заголов" in style_name.lower())
        if is_heading:
            headings.append(text)
        para_info.append((text, is_heading))

    low_heads = [h.lower() for h in headings]
    has_review = any("обзор литературы" in h or "литературный обзор" in h for h in low_heads)
    has_results = any("результат" in h for h in low_heads)
    has_discussion = any("обсуждение" in h or "вывод" in h for h in low_heads)

    text_sources = _estimate_sources_count(plain)
    url_count = _docx_bibliography_url_paragraph_count(doc)
    # Приоритет (handoff §dissertation_metrics, ветка с Танановой):
    # 1) Word-нумерация (numPr): auto-list, длина которого ≈ числу записей.
    #    python-docx не выводит auto-номера в paragraph.text, поэтому без
    #    XML мы её не «видим» (Гизатова 106, Досанов 40 — text-эвристика
    #    возвращала 29/4).
    #    КОРРЕКЦИЯ: если URL-абзацев в окне библиографии больше длины
    #    word-list (т.е. в библиографии есть записи без auto-нумерации),
    #    берём URL count. Реальный кейс: Тананова — word-list = 38, но
    #    реально 43 записи, каждая с ровно одним https:// (5 «висят» без
    #    numPr, обычно URL отдельной строкой). 43/38 = 1.13 — корректно.
    # 2) Реальная нумерация «1. … N.» / «[1] …» в начале абзацев библиографии:
    #    text-эвристика max(n.) надёжна (Сулейменова 45, Мараджапова 42).
    # 3) Подсчёт URL-абзацев в окне библиографии (fallback для случаев, когда
    #    нумерации нет вовсе): типовой ГОСТ — каждая запись содержит свой URL.
    #    Без этого fallback-а text-индексы ловили хвосты диапазонов страниц/
    #    годов и возвращали мусор.
    if word_list_count is not None and word_list_count >= 3:
        if url_count is not None and url_count > word_list_count:
            sources = url_count
        else:
            sources = word_list_count
    elif _docx_bibliography_has_line_numbering(doc):
        sources = text_sources
    elif url_count is not None:
        sources = url_count
    else:
        sources = text_sources
    formatting_compliance, font_size_ratio, font_family_ratio, line_spacing_ratio = (
        _analyze_docx_formatting(doc)
    )

    # pages_total: docProps/app.xml у DOCX, экспортированных из Google Docs
    # (или конвертированных некоторыми редакторами), часто содержит «<Pages>1</Pages>»
    # — счётчик, который Word-конвертер не пересчитал перед сохранением.
    # Реальный кейс (Камзебаева): plain ≈ 151k символов (≈68 стр.), но
    # docProps говорит 1. Безусловно доверять <Pages> нельзя.
    # Sanity-check: если оценка по символам в 5+ раз превышает значение из
    # метаданных, считаем метаданные занижёнными и используем оценку.
    chars_estimate = max(1, len(plain) // _CHARS_PER_PAGE_RU)
    pages_meta = _docx_page_count(docx_bytes)
    if pages_meta and pages_meta > 0 and chars_estimate >= 5 * pages_meta:
        pages_total = chars_estimate
    else:
        pages_total = pages_meta
    approx_pages = pages_total if pages_total and pages_total > 0 else chars_estimate

    # review section: find heading then accumulate until next heading
    review_text = ""
    start_idx = None
    for i, (t, is_h) in enumerate(para_info):
        if is_h and ("обзор литературы" in t.lower() or "литературный обзор" in t.lower()):
            start_idx = i + 1
            break
    if start_idx is not None:
        buf: list[str] = []
        for t, is_h in para_info[start_idx:]:
            if is_h:
                break
            buf.append(t)
        review_text = "\n".join(buf).strip()

    review_pages = None
    review_sources = None
    if review_text:
        # оценка страниц секции: пропорция по символам относительно общего текста и page_count
        if pages_total and pages_total > 0 and len(plain) > 0:
            ratio = len(review_text) / max(1, len(plain))
            review_pages = max(1, round(ratio * pages_total))
        else:
            review_pages = max(1, len(review_text) // _CHARS_PER_PAGE_RU)

        review_sources = _max_citation_index_in_text_chunk(review_text)

    notes: list[str] = []
    if pages_total is None:
        notes.append("Страницы в docx не найдены (docProps/app.xml).")
    if review_text and review_sources is None:
        notes.append("Источники в обзоре не оценены (не найден шаблон нумерации).")
    if formatting_compliance is None:
        notes.append("Соответствие оформлению не оценено (не хватило данных по стилям).")

    bib_heading_warn = bibliography_heading_issue_note(plain)
    if bib_heading_warn:
        notes.append(bib_heading_warn)

    margins_list = _docx_collect_section_margins(doc)
    page_margins, page_margins_secondary = _dominant_margins(margins_list)
    numbering_info = _docx_page_numbering_info(doc, docx_bytes)

    return DissertationMetrics(
        approx_pages=approx_pages,
        pdf_pages=None,
        sources_count=sources,
        review_pages=review_pages,
        review_sources_count=review_sources,
        has_literature_review=has_review,
        has_results=has_results,
        has_discussion=has_discussion,
        formatting_compliance=formatting_compliance,
        font_size_14_ratio=font_size_ratio,
        times_new_roman_ratio=font_family_ratio,
        single_spacing_ratio=line_spacing_ratio,
        headings_found=headings[:30],
        notes=notes,
        page_margins_cm=page_margins,
        page_margins_secondary_cm=page_margins_secondary,
        page_numbering_present=numbering_info.get("present"),
        page_numbering_position=numbering_info.get("position"),
        page_numbering_sections_with_footer=numbering_info.get("sections_with_footer"),
        page_numbering_sections_total=numbering_info.get("sections_total"),
        bibliography_heading_warning=bib_heading_warn,
    )


def _estimate_review_metrics(plain: str) -> tuple[int | None, int | None]:
    """
    Пытается выделить секцию «Обзор литературы» по plain text и оценить:
    - страницы секции (по _CHARS_PER_PAGE_RU)
    - число источников в секции (по нумерованным строкам)
    """
    low = plain.lower()
    start = low.find("обзор литературы")
    if start < 0:
        start = low.find("литературный обзор")
    if start < 0:
        return None, None
    # Конец: первое вхождение типичных следующих разделов после обзора
    end = len(plain)
    for marker in ("результат", "обсуждение", "заключение", "вывод", "глава 2", "глава ii"):
        pos = low.find(marker, start + 50)
        if 0 <= pos < end:
            end = pos
    seg = plain[start:end]
    pages = max(1, len(seg) // _CHARS_PER_PAGE_RU) if seg.strip() else None
    n = _max_citation_index_in_text_chunk(seg)
    return pages, n
