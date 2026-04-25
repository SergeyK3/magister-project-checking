"""Извлечение метаданных диссертации: тема (`dissertation_title`)
и язык (`dissertation_language`).

Логика разделена на «извлечение из Google Doc» и «извлечение из .docx-байтов»,
чтобы вызывающий код (report_enrichment, backfill-скрипт, бот) выбирал
подходящую ветку без дополнительной диспетчеризации.

Решения, согласованные в handoff §5 (2026-04-25):
- При неуверенной эвристике темы — возвращаем пустую строку, без issue.
- Язык записываем словами в нижнем регистре: «русский», «казахский», «английский».
- Английский язык фиксируем, но дополнительно поднимаем предупреждение через
  ``warn_if_unusual_language`` (логгер ``magister_checking.dissertation_meta``).
- Язык определяем по введению / первым ~5 страницам, не по всему телу,
  чтобы казахская/русская аннотация в начале/конце документа не «перетягивала»
  результат.
"""

from __future__ import annotations

import io
import logging
import re
from typing import Any

from docx import Document  # python-docx — уже зависимость через analyze_docx_bytes

from magister_checking.dissertation_metrics import iter_heading_texts
from magister_checking.docs_extract import extract_plain_text

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Константы языка
# ---------------------------------------------------------------------------

LANGUAGE_RUSSIAN = "русский"
LANGUAGE_KAZAKH = "казахский"
LANGUAGE_ENGLISH = "английский"

EXPECTED_LANGUAGES: frozenset[str] = frozenset({LANGUAGE_RUSSIAN, LANGUAGE_KAZAKH})
"""Языки, которые ожидаемы для магистерских диссертаций КазНУ.

Английский технически допустим, но требует подтверждения (см. handoff §5 п.4).
"""


# ---------------------------------------------------------------------------
# Тема диссертации
# ---------------------------------------------------------------------------

# Стоп-фильтр для строк, которые точно НЕ являются темой:
# 1) ``_STOP_PHRASES_SUBSTR`` — длинные характерные словосочетания
#    с титульного листа. Содержат уникальные сочетания, которые крайне
#    маловероятны в формулировке темы.
# 2) ``_STOP_PHRASES_EXACT`` — одиночные слова (заголовки разделов, отдельные
#    существительные). Проверяются на ТОЧНОЕ равенство, иначе валидные темы
#    вроде «Цифровая трансформация образования в Казахстане» отбрасывались бы
#    по подстроке «образован».
# 3) ``_LOCATION_LINE_RE`` — строки вида «Алматы 2026» внизу титула.
_STOP_PHRASES_SUBSTR: tuple[str, ...] = (
    "министерств",
    "магистерская диссертация",
    "магистрлік диссертация",
    "диссертация на соискание",
    "имени аль-фараби",
    "имени абая",
    "имени сатпаева",
    "научный руководитель",
    "ғылыми жетекші",
)

_STOP_PHRASES_EXACT: frozenset[str] = frozenset(
    {
        "содержание",
        "оглавление",
        "мазмұны",
        "введение",
        "кіріспе",
        "abstract",
        "annotation",
        "аннотация",
        "аңдатпа",
        "университет",
        "факультет",
        "кафедра",
        "институт",
        "образование",
        "нормативные ссылки",
        "нормативтік сілтемелер",
        "определения",
        "анықтамалар",
        "обозначения и сокращения",
        "белгілеулер мен қысқартулар",
        "белгілер мен қысқартулар",
        "перечень условных обозначений",
        "перечень сокращений",
        "список сокращений",
        "список рисунков",
        "список таблиц",
        "кестелер мен тізімдер",
        "обзор литературы",
        "литературный обзор",
        "методология",
        "материалы и методы",
        "методы исследования",
        "результаты",
        "результаты исследования",
        "обсуждение",
        "заключение",
        "қорытынды",
        "выводы",
        "приложение",
        "приложения",
        "глоссарий",
        "список использованных источников",
        "список использованной литературы",
        "пайдаланылған әдебиеттер тізімі",
        "references",
    }
)

# Подстрочные стоп-паттерны для нумерованных глав и подразделов («Глава 1»,
# «Раздел 2», «1 Обзор литературы», «1.1 Теоретические основы»). Эти строки
# никогда не являются темой; ужесточаем фильтр, чтобы Heading-fallback не
# подцепил их.
_STOP_NUMBERED_SECTION_RE = re.compile(
    r"(?im)^\s*(?:глава|раздел|часть|тарау|бөлім|chapter|section)\s+\d|"
    r"^\s*\d+\s*\.?\s*(?:обзор\s+литератур|введени|заключени|"
    r"теоретическ|методолог|материал|анализ|обсуждени|результат)|"
    r"^\s*\d+\.\d+"
)

_LOCATION_LINE_RE = re.compile(
    r"(?i)^(алматы|астана|нур[-\s]?султан|шымкент|караганда|тараз|павлодар)"
    r"\s*\d{0,4}\s*$"
)

# Библиографические префиксы, которые иногда стоят CAPS-строкой на титуле:
# «УДК: 614.2», «МПК: G16H 20/00», «ББК 65», «ISBN ...», «ISSN ...».
_BIBLIO_PREFIX_RE = re.compile(r"(?i)^\s*(удк|мпк|ббк|isbn|issn|udc|doi)[\s:№#]")

# ФИО автора, написанное CAPS: «СУЛЕЙМЕНОВА ИНДИРА САРСЕНБЕКОВНА». Слов
# 2..4, и хотя бы одно с типичным русским / казахским ФИО-суффиксом.
_FIO_SUFFIX_RE = re.compile(
    r"(?i)(?:"
    r"ов|ова|ев|ева|ёв|ёва|ин|ина|ын|ына|"
    r"ович|овна|евич|евна|"
    r"оглы|оғлы|қызы|ұлы|улы|"
    r"овский|евский|инский|"
    r"енко|чук|юк|ук"
    r")$"
)

# «На тему: <тема>» в одной строке — inline-маркер для шаблонов, где он
# присутствует. Двоеточие или тире обязательны, иначе словосочетание «на тему»
# в произвольном предложении ложно срабатывало бы.
_ON_TOPIC_INLINE_RE = re.compile(
    r"(?im)\bна\s+тему\s*[:\-\u2013\u2014]\s*[«\"„\u00ab\u201c\u201e]?\s*(.+)$",
)

# Параграф, состоящий ТОЛЬКО из маркера «На тему» / «На тему:». Тема —
# в следующем непустом параграфе (типичный вариант, когда титул сверстан
# двумя строками).
_ON_TOPIC_HEADER_RE = re.compile(
    r"(?im)^\s*на\s+тему\s*[:\-\u2013\u2014]?\s*[«\"„\u00ab\u201c\u201e]?\s*$",
)

# «Тема [магистерской] диссертации: …»
_TOPIC_HEADER_RE = re.compile(
    r"(?im)\bтема\s+(?:магистерской\s+)?диссертации\s*[:\-\u2013\u2014]\s*"
    r"[«\"„\u00ab\u201c\u201e]?\s*(.+)$",
)

# Маркер «это магистерская работа», который встречается в первых ~10
# параграфах ГОСО-титула. Используется как «якорь»: тема — параграф ВЫШЕ него.
# Подтверждено на двух реальных диссертациях (Камзебаева, Сапарбаева,
# probe от 2026-04-25).
_DEGREE_MARKER_RE = re.compile(
    r"(?im)\b("
    r"магистерск(?:ая|ий)\s+(?:диссертац|проект)"
    r"|магистрл[іi]к\s+(?:жоба|диссертац)"
    r"|магистр\s+(?:академиялы[қk]\s+)?[дd][әa]режес"
    r"|соискани[ея]\s+степени\s+магистр"
    r"|степени\s+магистр[а]?\s+(?:здравоохранения|общественного|техническ|социальн|педагогик|экономик|деловог)"
    r"|денсаулы[қk]\s+са[қk]тау\s+магистр"
    r")"
)

_QUOTE_CHARS = "«»\"„""'\u00ab\u00bb\u201c\u201d\u201e\u201f\u2018\u2019"

_MIN_TITLE_LEN = 5
_MAX_TITLE_LEN = 300


def _clean_title(raw: str) -> str:
    text = (raw or "").strip()
    text = re.sub(r"\s+", " ", text)
    text = text.strip(_QUOTE_CHARS + " .,;:")
    text = re.sub(r"\s+", " ", text).strip()
    return text


def _is_stop_phrase(text: str) -> bool:
    raw = text or ""
    low = raw.lower()
    norm = re.sub(r"\s+", " ", low.strip(_QUOTE_CHARS + " .,;:"))
    if not norm:
        return True
    if norm in _STOP_PHRASES_EXACT:
        return True
    if _LOCATION_LINE_RE.match(norm):
        return True
    if _STOP_NUMBERED_SECTION_RE.search(norm):
        return True
    if _BIBLIO_PREFIX_RE.match(raw):
        return True
    if _looks_like_fio(raw):
        return True
    return any(stop in norm for stop in _STOP_PHRASES_SUBSTR)


def _looks_like_fio(text: str) -> bool:
    """ФИО автора (2-4 слова, каждое с заглавной, хотя бы одно с ФИО-суффиксом).

    Требование «каждое слово начинается с заглавной» отсекает родительные
    падежи в темах вроде «Оптимизация алгоритмов сортировки», где
    «алгоритмов» заканчивается на «-ов», но не является фамилией.

    Покрывает русские ФИО («Сулейменова Индира Сарсенбековна»), казахские
    («Камзебаева Анель Дулатовна», «Сапарбаева Жайна Саматқызы») и CAPS-форму
    («СУЛЕЙМЕНОВА ИНДИРА САРСЕНБЕКОВНА»).
    """

    words = [w for w in re.split(r"\s+", text.strip()) if w]
    if not (2 <= len(words) <= 4):
        return False
    cleaned = [w.strip(_QUOTE_CHARS + ".,;:()") for w in words]
    for w in cleaned:
        if not w:
            return False
        first = w[0]
        if not first.isalpha() or not first.isupper():
            return False
    return any(_FIO_SUFFIX_RE.search(w) for w in cleaned)


def _is_caps_paragraph(text: str, *, ratio_threshold: float = 0.7) -> bool:
    """True, если ≥ ``ratio_threshold`` буквенных символов — заглавные.

    ФИО автора в типичном титуле выглядит «Фамилия Имя Отчество» — там
    всего 3 заглавные буквы из ~25 → ratio ≈ 0.12. Тема же оформлена
    «СПЛОШНЫМ КАПСОМ» → ratio ≈ 1.0.
    """

    letters = [ch for ch in text if ch.isalpha()]
    if not letters:
        return False
    upper = sum(1 for ch in letters if ch.isupper())
    return (upper / len(letters)) >= ratio_threshold


def _detect_title_from_govt_template(paragraphs: list[str]) -> str:
    """Извлечь тему из ГОСО-титула КазНУ/АМУ.

    Алгоритм (probe от 2026-04-25):
    1. В первых ~30 параграфах ищем «магистерский проект / магистрлік жоба /
       соискание степени магистра». Это надёжный «якорь» — он в типичном
       титуле находится сразу после темы и кода специальности.
    2. От якоря идём ВВЕРХ (на 1..6 параграфов) и берём первый CAPS-параграф
       длиной 20..300 символов, не из стоп-списка. ФИО автора отбрасывается
       по CAPS-проверке (ФИО — «Иванов И. И.», не CAPS), служебные
       строки — по стоп-фильтру.
    """

    head_window = paragraphs[:30]
    for idx, paragraph in enumerate(head_window):
        if not _DEGREE_MARKER_RE.search(paragraph):
            continue
        for k in range(idx - 1, max(idx - 7, -1), -1):
            candidate = paragraphs[k].strip()
            if not candidate:
                continue
            if not (20 <= len(candidate) <= 300):
                continue
            if _is_stop_phrase(candidate):
                continue
            if not _is_caps_paragraph(candidate):
                continue
            cleaned = _clean_title(candidate)
            if _looks_like_title(cleaned):
                return cleaned
        # Якорь нашли, но темы над ним нет — дальше ловить нечего, выходим
        break
    return ""


def _looks_like_title(text: str) -> bool:
    if not text:
        return False
    if not (_MIN_TITLE_LEN <= len(text) <= _MAX_TITLE_LEN):
        return False
    if _is_stop_phrase(text):
        return False
    # Темой не может быть только число / только знаки препинания.
    if not re.search(r"[A-Za-zА-Яа-яёЁ\u0400-\u04FF]", text):
        return False
    return True


def _first_paragraphs_from_plain(text: str, *, limit: int = 80) -> list[str]:
    """Первые ``limit`` непустых строк ``plain_text``.

    Полагаемся на то, что Google Docs / python-docx разделяют параграфы
    переводом строки, поэтому ``split('\\n')`` достаточно для нашей эвристики.
    """

    out: list[str] = []
    for raw in (text or "").split("\n"):
        s = raw.strip()
        if not s:
            continue
        out.append(s)
        if len(out) >= limit:
            break
    return out


def _detect_title_from_paragraphs(paragraphs: list[str], headings: list[str]) -> str:
    """Применяет последовательность эвристик к подготовленным данным.

    Порядок эвристик от самой надёжной к наименее:

    1. ГОСО-титул: маркер «магистерский проект / магистрлік жоба» как
       якорь, тема — CAPS-параграф над ним. Подтверждено на двух реальных
       диссертациях (probe 2026-04-25).
    2. «На тему: …» в одной строке.
    3. «На тему:» отдельной строкой → тема в следующем непустом параграфе.
    4. «Тема [магистерской] диссертации: …».
    5. Первый Heading 1, не из стоп-списка.

    Если ни один шаг не сработал — возвращаем "" (handoff §5 п.3 = empty).
    """

    govt_title = _detect_title_from_govt_template(paragraphs)
    if govt_title:
        return govt_title

    head_window = paragraphs[:60]

    for paragraph in head_window:
        match = _ON_TOPIC_INLINE_RE.search(paragraph)
        if not match:
            continue
        cleaned = _clean_title(match.group(1) or "")
        if _looks_like_title(cleaned):
            return cleaned

    for idx, paragraph in enumerate(head_window):
        if not _ON_TOPIC_HEADER_RE.match(paragraph):
            continue
        for next_paragraph in head_window[idx + 1 : idx + 4]:
            cleaned = _clean_title(next_paragraph)
            if _looks_like_title(cleaned):
                return cleaned

    for paragraph in head_window:
        match = _TOPIC_HEADER_RE.search(paragraph)
        if not match:
            continue
        cleaned = _clean_title(match.group(1) or "")
        if _looks_like_title(cleaned):
            return cleaned

    for heading in headings[:10]:
        cleaned = _clean_title(heading)
        if _looks_like_title(cleaned):
            return cleaned

    return ""


def detect_dissertation_title_from_gdoc(document: dict[str, Any]) -> str:
    """Извлекает тему диссертации из Google Doc (`documents.get` JSON)."""

    plain = extract_plain_text(document) if document else ""
    paragraphs = _first_paragraphs_from_plain(plain, limit=80)
    headings = [h.strip() for h in iter_heading_texts(document or {}) if (h or "").strip()]
    return _detect_title_from_paragraphs(paragraphs, headings)


def detect_dissertation_title_from_docx_bytes(blob: bytes) -> str:
    """Извлекает тему диссертации из .docx (байты файла)."""

    if not blob:
        return ""
    try:
        doc = Document(io.BytesIO(blob))
    except Exception as exc:  # noqa: BLE001
        logger.warning("dissertation_meta: не удалось открыть .docx: %s", exc)
        return ""

    paragraphs: list[str] = []
    headings: list[str] = []
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if not text:
            continue
        paragraphs.append(text)
        try:
            style_name = str(para.style.name or "")
        except Exception:  # noqa: BLE001
            style_name = ""
        if re.search(r"(?i)\bheading\b", style_name) or "заголов" in style_name.lower():
            headings.append(text)
        if len(paragraphs) >= 80:
            break

    return _detect_title_from_paragraphs(paragraphs, headings)


# ---------------------------------------------------------------------------
# Язык диссертации
# ---------------------------------------------------------------------------

# Кириллические буквы, специфичные для казахского алфавита. Их доля
# относительно всех кириллических букв — надёжный признак рус/каз.
_KAZAKH_SPECIFIC_LETTERS = "әңғқөұүһіӘҢҒҚӨҰҮҺІ"

# Доля латиницы в тексте, выше которой считаем диссертацию англоязычной.
_ENGLISH_LATIN_RATIO_THRESHOLD = 0.6

# Доля казахских специфичных букв среди кириллицы, выше которой → казахский.
# В чистом казахском тексте обычно 5–10%; в русском «і»/«ң» практически
# не встречаются (только в именах собственных). 2% — безопасный порог.
_KAZAKH_RATIO_THRESHOLD = 0.02

# Маркеры начала «основного» текста — после них берём 5000 символов для языка.
_INTRO_HEADER_RE = re.compile(
    r"(?im)^\s*(введение|кіріспе|kіrіspе|introduction)\s*$",
)

_LANGUAGE_SAMPLE_LEN = 5000
_TITLE_PAGE_SKIP_CHARS = 500
_LANGUAGE_SCAN_WINDOW = 60000


def _slice_for_language(plain_text: str) -> str:
    """Подбирает фрагмент для подсчёта статистики языка.

    1. Если в первых ``_LANGUAGE_SCAN_WINDOW`` символах есть отдельная строка
       «ВВЕДЕНИЕ» / «КІРІСПЕ» / «INTRODUCTION» — берём ``_LANGUAGE_SAMPLE_LEN``
       символов после неё.
    2. Иначе — берём срез ``[_TITLE_PAGE_SKIP_CHARS : _TITLE_PAGE_SKIP_CHARS + _LANGUAGE_SAMPLE_LEN]``,
       чтобы пропустить университетский титул, который у всех одинаков.
    3. Если документ короче, чем требуемый сдвиг — возвращаем весь текст.
    """

    if not plain_text:
        return ""
    snapshot = plain_text[:_LANGUAGE_SCAN_WINDOW]
    intro_match = _INTRO_HEADER_RE.search(snapshot)
    if intro_match:
        start = intro_match.end()
        return snapshot[start : start + _LANGUAGE_SAMPLE_LEN]
    if len(snapshot) <= _TITLE_PAGE_SKIP_CHARS + _LANGUAGE_SAMPLE_LEN:
        return snapshot
    return snapshot[_TITLE_PAGE_SKIP_CHARS : _TITLE_PAGE_SKIP_CHARS + _LANGUAGE_SAMPLE_LEN]


def detect_dissertation_language_from_text(text: str) -> str:
    """«русский» / «казахский» / «английский» / «» (если букв вообще нет)."""

    chunk = _slice_for_language(text)
    if not chunk:
        return ""

    cyrillic = 0
    latin = 0
    kazakh_specific = 0
    for ch in chunk:
        if ("а" <= ch <= "я") or ("А" <= ch <= "Я") or ch in "ёЁ":
            cyrillic += 1
        if ch in _KAZAKH_SPECIFIC_LETTERS:
            kazakh_specific += 1
            cyrillic += 1
        if ("a" <= ch <= "z") or ("A" <= ch <= "Z"):
            latin += 1

    total_letters = cyrillic + latin
    if total_letters == 0:
        return ""

    if latin / total_letters > _ENGLISH_LATIN_RATIO_THRESHOLD:
        return LANGUAGE_ENGLISH

    if cyrillic == 0:
        return ""

    if kazakh_specific / cyrillic > _KAZAKH_RATIO_THRESHOLD:
        return LANGUAGE_KAZAKH
    return LANGUAGE_RUSSIAN


def detect_dissertation_language_from_gdoc(document: dict[str, Any]) -> str:
    if not document:
        return ""
    plain = extract_plain_text(document)
    return detect_dissertation_language_from_text(plain)


def detect_dissertation_language_from_docx_bytes(blob: bytes) -> str:
    if not blob:
        return ""
    try:
        doc = Document(io.BytesIO(blob))
    except Exception as exc:  # noqa: BLE001
        logger.warning("dissertation_meta: не удалось открыть .docx: %s", exc)
        return ""
    plain = "\n".join((p.text or "") for p in doc.paragraphs)
    return detect_dissertation_language_from_text(plain)


def warn_if_unusual_language(language: str, *, context: str = "") -> None:
    """Логирует предупреждение, если язык не из ``EXPECTED_LANGUAGES``.

    Сейчас это только английский (handoff §5 п.4 «warn»). Решение об остановке
    pipeline / показе магистранту делается выше по стеку — здесь лишь логируем,
    чтобы запись в Sheets и дальнейший прогон не блокировались.
    """

    if not language or language in EXPECTED_LANGUAGES:
        return
    logger.warning(
        "dissertation_meta: неожиданный язык диссертации %r (%s)",
        language,
        context or "контекст не указан",
    )
