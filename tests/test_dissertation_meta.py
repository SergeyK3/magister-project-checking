"""Тесты эвристик определения темы и языка диссертации.

Покрытие соответствует решениям handoff §5 (2026-04-25):
- Тема: «На тему: …», «Тема диссертации: …», fallback на Heading 1, иначе "".
- Язык: «русский» / «казахский» / «английский»; пустая строка — если букв нет.
- Английский язык логирует предупреждение; русский/казахский — нет.
"""

from __future__ import annotations

import io
import logging
import unittest

from docx import Document  # type: ignore[import-untyped]

from magister_checking.dissertation_meta import (
    LANGUAGE_ENGLISH,
    LANGUAGE_KAZAKH,
    LANGUAGE_RUSSIAN,
    _normalize_kazakh_text,
    detect_dissertation_language_from_docx_bytes,
    detect_dissertation_language_from_gdoc,
    detect_dissertation_language_from_text,
    detect_dissertation_title_from_docx_bytes,
    detect_dissertation_title_from_gdoc,
    warn_if_unusual_language,
)


# ---------------------------------------------------------------------------
# Хелперы для синтетических Google Docs JSON
# ---------------------------------------------------------------------------


def _paragraph(text: str, *, heading_level: int | None = None) -> dict:
    named_style = "NORMAL_TEXT" if heading_level is None else f"HEADING_{heading_level}"
    return {
        "paragraph": {
            "elements": [{"textRun": {"content": text, "textStyle": {}}}],
            "paragraphStyle": {"namedStyleType": named_style},
        }
    }


def _document(content: list[dict]) -> dict:
    return {"body": {"content": content}}


def _build_docx_bytes(paragraphs: list[tuple[str, str | None]]) -> bytes:
    """Собирает .docx в памяти из списка ``(text, style_name)``.

    ``style_name`` = ``"Heading 1"`` / ``"Heading 2"`` сделает параграф
    заголовком соответствующего уровня (через ``add_heading``, чтобы
    стиль гарантированно существовал в свежесозданном документе).
    """

    doc = Document()
    for text, style in paragraphs:
        if style and style.lower().startswith("heading"):
            try:
                level = int(style.split()[-1])
            except ValueError:
                level = 1
            doc.add_heading(text, level=level)
        else:
            doc.add_paragraph(text)
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ---------------------------------------------------------------------------
# Тема диссертации — Google Doc
# ---------------------------------------------------------------------------


class DetectTitleFromGdocTests(unittest.TestCase):
    def test_on_topic_inline(self) -> None:
        document = _document(
            [
                _paragraph("Министерство образования и науки\n"),
                _paragraph("Магистерская диссертация\n"),
                _paragraph("На тему: Оптимизация алгоритмов сортировки\n"),
                _paragraph("Алматы 2026\n"),
            ]
        )
        self.assertEqual(
            detect_dissertation_title_from_gdoc(document),
            "Оптимизация алгоритмов сортировки",
        )

    def test_on_topic_next_paragraph(self) -> None:
        document = _document(
            [
                _paragraph("Магистерская диссертация\n"),
                _paragraph("На тему:\n"),
                _paragraph("Применение нейросетей в медицинской диагностике\n"),
                _paragraph("Выполнил: Иванов И. И.\n"),
            ]
        )
        self.assertEqual(
            detect_dissertation_title_from_gdoc(document),
            "Применение нейросетей в медицинской диагностике",
        )

    def test_topic_header_pattern(self) -> None:
        document = _document(
            [
                _paragraph("Магистерская диссертация\n"),
                _paragraph(
                    "Тема магистерской диссертации: Анализ больших данных в банковском секторе\n"
                ),
            ]
        )
        self.assertEqual(
            detect_dissertation_title_from_gdoc(document),
            "Анализ больших данных в банковском секторе",
        )

    def test_strips_quotes(self) -> None:
        document = _document(
            [_paragraph("На тему: «Машинное обучение и его применения»\n")]
        )
        self.assertEqual(
            detect_dissertation_title_from_gdoc(document),
            "Машинное обучение и его применения",
        )

    def test_falls_back_to_heading_when_no_marker(self) -> None:
        document = _document(
            [
                _paragraph("Содержание\n"),
                _paragraph(
                    "Разработка веб-приложения для учёта успеваемости\n",
                    heading_level=1,
                ),
                _paragraph("Введение\n", heading_level=1),
                _paragraph("Глава 1\n", heading_level=1),
            ]
        )
        self.assertEqual(
            detect_dissertation_title_from_gdoc(document),
            "Разработка веб-приложения для учёта успеваемости",
        )

    def test_skips_stop_phrases_in_headings(self) -> None:
        document = _document(
            [
                _paragraph("Содержание\n", heading_level=1),
                _paragraph("Введение\n", heading_level=1),
                _paragraph("Магистерская диссертация\n", heading_level=1),
            ]
        )
        self.assertEqual(detect_dissertation_title_from_gdoc(document), "")

    def test_returns_empty_for_empty_doc(self) -> None:
        self.assertEqual(detect_dissertation_title_from_gdoc({}), "")
        self.assertEqual(detect_dissertation_title_from_gdoc(_document([])), "")

    def test_ignores_too_short_capture(self) -> None:
        document = _document([_paragraph("На тему: ОК\n")])
        self.assertEqual(detect_dissertation_title_from_gdoc(document), "")

    def test_ignores_capture_without_letters(self) -> None:
        document = _document([_paragraph("На тему: 12345\n")])
        self.assertEqual(detect_dissertation_title_from_gdoc(document), "")

    def test_govt_template_caps_title_above_degree_marker_ru(self) -> None:
        # Структура реальной диссертации Сапарбаевой (probe 2026-04-25).
        document = _document(
            [
                _paragraph("НАО «Медицинский университет Астана»\n"),
                _paragraph("УДК: 614.2:005.96:004.738.5\n"),
                _paragraph("МПК: G06Q 10/06\n"),
                _paragraph("Сапарбаева Жайна Саматқызы\n"),
                _paragraph(
                    "СИСТЕМАТИЗАЦИЯ В КОРПОРАТИВНОЙ БИБЛИОТЕКЕ "
                    "ДОЛЖНОСТНЫХ ИНСТРУКЦИЙ МЕДИЦИНСКОЙ ОРГАНИЗАЦИИ\n"
                ),
                _paragraph("7M10116 – «Общественное здравоохранение»\n"),
                _paragraph(
                    "Магистерский проекта на соискание степени магистра здравоохранения\n"
                ),
                _paragraph("Научный руководитель: д.м.н., профессор\n"),
                _paragraph("Астана 2026 г\n"),
            ]
        )
        self.assertEqual(
            detect_dissertation_title_from_gdoc(document),
            "СИСТЕМАТИЗАЦИЯ В КОРПОРАТИВНОЙ БИБЛИОТЕКЕ "
            "ДОЛЖНОСТНЫХ ИНСТРУКЦИЙ МЕДИЦИНСКОЙ ОРГАНИЗАЦИИ",
        )

    def test_govt_template_academic_degree_marker_with_extra_word(self) -> None:
        document = _document(
            [
                _paragraph("НАО «Медицинский университет Астана»\n"),
                _paragraph("СЕРИКОВА ТАЛШЫН КАЙРГАЛИЕВНА\n"),
                _paragraph(
                    "ОПТИМИЗАЦИЯ АМБУЛАТОРНОЙ МЕДИЦИНСКОЙ ПОМОЩИ БОЛЬНЫМ "
                    "С РЕВМАТОЛОГИЧЕСКИМИ ЗАБОЛЕВАНИЯМИ\n"
                ),
                _paragraph("7M10105 – «Менеджмент в здравоохранении»\n"),
                _paragraph(
                    "Диссертация на соискание академической степени магистра "
                    "медицинских наук\n"
                ),
                _paragraph("Актуальность темы\n", heading_level=1),
            ]
        )
        self.assertEqual(
            detect_dissertation_title_from_gdoc(document),
            "ОПТИМИЗАЦИЯ АМБУЛАТОРНОЙ МЕДИЦИНСКОЙ ПОМОЩИ БОЛЬНЫМ "
            "С РЕВМАТОЛОГИЧЕСКИМИ ЗАБОЛЕВАНИЯМИ",
        )

    def test_govt_template_joins_wrapped_caps_title_above_split_degree_marker(self) -> None:
        document = _document(
            [
                _paragraph("НАО «Медицинский университет Астана»\n"),
                _paragraph("Ибраев Руслан Тургумбекович\n"),
                _paragraph(
                    "АНАЛИЗ УДОВЛЕТВОРЁННОСТИ ВЗРОСЛОГО СЕЛЬСКОГО НАСЕЛЕНИЯ "
                    "ДОСТУПНОСТЬЮ ЛЕКАРСТВЕННЫХ СРЕДСТВ (НА ПРИМЕРЕ АКМОЛИНСКОЙ\n"
                ),
                _paragraph("ОБЛАСТИ)\n"),
                _paragraph("7M10116 – «Общественное здравоохранение»\n"),
                _paragraph("Проект на соискание академической\n"),
                _paragraph("степени магистра здравоохранения\n"),
            ]
        )
        self.assertEqual(
            detect_dissertation_title_from_gdoc(document),
            "АНАЛИЗ УДОВЛЕТВОРЁННОСТИ ВЗРОСЛОГО СЕЛЬСКОГО НАСЕЛЕНИЯ "
            "ДОСТУПНОСТЬЮ ЛЕКАРСТВЕННЫХ СРЕДСТВ (НА ПРИМЕРЕ АКМОЛИНСКОЙ ОБЛАСТИ)",
        )

    def test_govt_template_accepts_quoted_mixed_case_title_above_degree_marker(self) -> None:
        document = _document(
            [
                _paragraph("НАО «Медицинский университет Астана»\n"),
                _paragraph("Антикеева Алия Алмасовна\n"),
                _paragraph(
                    "«Научное обоснование комплекса мероприятий по развитию "
                    "посмертного донорства в Республике Казахстан»\n"
                ),
                _paragraph("7М10105 – «Менеджмент в здравоохранении»\n"),
                _paragraph(
                    "Диссертация на соискание академической степени магистра "
                    "медицинских наук\n"
                ),
            ]
        )
        self.assertEqual(
            detect_dissertation_title_from_gdoc(document),
            "Научное обоснование комплекса мероприятий по развитию "
            "посмертного донорства в Республике Казахстан",
        )

    def test_govt_template_accepts_plain_mixed_case_title_between_fio_and_specialty(self) -> None:
        # Реальный шаблон row 21: тема не CAPS и не в кавычках, но находится
        # на титуле между ФИО автора и шифром специальности.
        document = _document(
            [
                _paragraph("НАО «Медицинский университет Астана»\n"),
                _paragraph("УДК: 614.2:005:339.138:004\n"),
                _paragraph("МПК:\n"),
                _paragraph("Джумабекова Инкар Сериковна\n"),
                _paragraph(
                    "Развитие системы цифрового маркетинга в управлении "
                    "медицинской организации\n"
                ),
                _paragraph("7М10121 «Менеджмент в здравоохранении»\n"),
                _paragraph(
                    "Магистерский проект на соискание степени магистра "
                    "здравоохранения\n"
                ),
            ]
        )
        self.assertEqual(
            detect_dissertation_title_from_gdoc(document),
            "Развитие системы цифрового маркетинга в управлении "
            "медицинской организации",
        )

    def test_govt_template_accepts_wrapped_mixed_case_kazakh_title_before_degree_marker(self) -> None:
        document = _document(
            [
                _paragraph("«Астана Медицина Университеті» КеАҚ\n"),
                _paragraph("УДК: 614.2\n"),
                _paragraph("МПК: A61B 5/12\n"),
                _paragraph("Қожахметова Аружан Серікқызы\n"),
                _paragraph("денсаулық сақтау ұйымындағы цифрлық\n"),
                _paragraph("процестерді жетілдіру\n"),
                _paragraph(
                    "магистрлік жоба (бейіндік бағыт үшін)\n"
                ),
            ]
        )
        self.assertEqual(
            detect_dissertation_title_from_gdoc(document),
            "денсаулық сақтау ұйымындағы цифрлық процестерді жетілдіру",
        )

    def test_govt_template_caps_title_above_degree_marker_kk(self) -> None:
        # Структура реальной диссертации Камзебаевой.
        document = _document(
            [
                _paragraph("«Астана Медицина Университеті» КеАҚ\n"),
                _paragraph("УДК: 613.644:656.2\n"),
                _paragraph("МПК: А61В 5/12; G16H 50/30\n"),
                _paragraph("Камзебаева Анель Дулатовна\n"),
                _paragraph("ӨҢДІРІСТІК ШУДЫҢ МАШИНИСТЕРДІҢ ДЕНСАУЛЫҒЫНА ӘСЕРІ\n"),
                _paragraph("7М10116 - \"Қоғамдық денсаулық сақтау\"\n"),
                _paragraph(
                    "денсаулық сақтау магистрі академиялық дәрежесін алуға арналған "
                    "магистрлік жоба (бейіндік бағыт үшін)\n"
                ),
                _paragraph("Астана 2026 ж.\n"),
            ]
        )
        self.assertEqual(
            detect_dissertation_title_from_gdoc(document),
            "ӨҢДІРІСТІК ШУДЫҢ МАШИНИСТЕРДІҢ ДЕНСАУЛЫҒЫНА ӘСЕРІ",
        )

    def test_govt_template_skips_author_fio_above_degree_marker(self) -> None:
        # ФИО автора над якорем — не CAPS, не должно подцепиться как тема.
        # Над ФИО других CAPS-параграфов нет → тема пустая.
        document = _document(
            [
                _paragraph("«Астана Медицина Университеті» КеАҚ\n"),
                _paragraph("Иванов Иван Иванович\n"),
                _paragraph(
                    "Магистерский проект на соискание степени магистра здравоохранения\n"
                ),
            ]
        )
        self.assertEqual(detect_dissertation_title_from_gdoc(document), "")

    def test_govt_template_skips_caps_fio_above_marker(self) -> None:
        # Реальный кейс row 6 Сулейменова: ФИО автора над маркером написано
        # CAPS, без отдельной строки темы. Выбирать ФИО как тему нельзя.
        document = _document(
            [
                _paragraph("НАО «Медицинский университет Астана»\n"),
                _paragraph("УДК: 614.2\n"),
                _paragraph("МПК: G16H 50/30\n"),
                _paragraph("СУЛЕЙМЕНОВА ИНДИРА САРСЕНБЕКОВНА\n"),
                _paragraph(
                    "Магистерский проект на соискание степени магистра здравоохранения\n"
                ),
            ]
        )
        self.assertEqual(detect_dissertation_title_from_gdoc(document), "")

    def test_govt_template_skips_classifier_above_marker(self) -> None:
        # Реальный кейс row 13 Тананова: над маркером — строка
        # «МПК: G16H 20/00, G16H 10/00, A61B 5/02» (классификатор, не тема).
        document = _document(
            [
                _paragraph("НАО «Медицинский университет Астана»\n"),
                _paragraph("Тананова Айнур Ахметовна\n"),
                _paragraph("МПК: G16H 20/00, G16H 10/00, A61B 5/02\n"),
                _paragraph(
                    "Магистерский проект на соискание степени магистра здравоохранения\n"
                ),
            ]
        )
        self.assertEqual(detect_dissertation_title_from_gdoc(document), "")

    def test_govt_template_skips_section_headings(self) -> None:
        # «НОРМАТИВНЫЕ ССЫЛКИ» / «ОПРЕДЕЛЕНИЯ» / «1 ОБЗОР ЛИТЕРАТУРЫ» —
        # стоп-фразы; даже если они в headings, тема не выбирается.
        document = _document(
            [
                _paragraph("Содержание\n", heading_level=1),
                _paragraph("НОРМАТИВНЫЕ ССЫЛКИ\n", heading_level=1),
                _paragraph("ОПРЕДЕЛЕНИЯ\n", heading_level=1),
                _paragraph("ОБОЗНАЧЕНИЯ И СОКРАЩЕНИЯ\n", heading_level=1),
                _paragraph("ВВЕДЕНИЕ\n", heading_level=1),
                _paragraph("1. ОБЗОР ЛИТЕРАТУРЫ\n", heading_level=1),
                _paragraph("1.1. Теоретические основы\n", heading_level=1),
            ]
        )
        self.assertEqual(detect_dissertation_title_from_gdoc(document), "")

    def test_heading_fallback_skips_service_headings_from_doc_without_title_page(self) -> None:
        # Реальный класс ошибок: документ начинается не с титульного листа, а со
        # служебных разделов. Их нельзя принимать за название диссертации.
        document = _document(
            [
                _paragraph("НОРМАТИВНЫЕ ССЫЛКИ\n", heading_level=1),
                _paragraph("ОПРЕДЕЛЕНИЯ\n", heading_level=1),
                _paragraph("ПЕРЕЧЕНЬ СОКРАЩЕНИЙ И ОБОЗНАЧЕНИЙ\n", heading_level=1),
                _paragraph("СПИСОК ТАБЛИЦ И РИСУНКОВ\n", heading_level=1),
                _paragraph("ВВЕДЕНИЕ\n", heading_level=1),
                _paragraph("Актуальность\n", heading_level=1),
                _paragraph("Цель исследования\n", heading_level=1),
                _paragraph("Объект и предмет исследования\n", heading_level=1),
                _paragraph("Практическая значимость\n", heading_level=1),
                _paragraph("Структура работы\n", heading_level=1),
            ]
        )
        self.assertEqual(detect_dissertation_title_from_gdoc(document), "")


# ---------------------------------------------------------------------------
# Тема диссертации — .docx
# ---------------------------------------------------------------------------


class DetectTitleFromDocxTests(unittest.TestCase):
    def test_on_topic_inline_in_docx(self) -> None:
        blob = _build_docx_bytes(
            [
                ("Министерство образования и науки", None),
                ("Магистерская диссертация", None),
                ("На тему: Распознавание речи на казахском языке", None),
                ("Алматы 2026", None),
            ]
        )
        self.assertEqual(
            detect_dissertation_title_from_docx_bytes(blob),
            "Распознавание речи на казахском языке",
        )

    def test_heading_fallback_in_docx(self) -> None:
        blob = _build_docx_bytes(
            [
                ("Содержание", None),
                ("Цифровая трансформация образования в Казахстане", "Heading 1"),
                ("Введение", "Heading 1"),
            ]
        )
        self.assertEqual(
            detect_dissertation_title_from_docx_bytes(blob),
            "Цифровая трансформация образования в Казахстане",
        )

    def test_govt_template_plain_mixed_case_title_in_docx(self) -> None:
        blob = _build_docx_bytes(
            [
                ("НАО «Медицинский университет Астана»", None),
                ("УДК: 614.2:005:339.138:004", None),
                ("МПК:", None),
                ("Джумабекова Инкар Сериковна", None),
                (
                    "Развитие системы цифрового маркетинга в управлении "
                    "медицинской организации",
                    None,
                ),
                ("7М10121 «Менеджмент в здравоохранении»", None),
                (
                    "Магистерский проект на соискание степени магистра "
                    "здравоохранения",
                    None,
                ),
            ]
        )
        self.assertEqual(
            detect_dissertation_title_from_docx_bytes(blob),
            "Развитие системы цифрового маркетинга в управлении "
            "медицинской организации",
        )

    def test_govt_template_mixed_case_title_before_degree_line_in_docx(self) -> None:
        blob = _build_docx_bytes(
            [
                ("НАО «Медицинский университет Астана»", None),
                ("Школа общественного здоровья", None),
                ("УДК: 614.2:658", None),
                ("МПК: A61B, G06Q10", None),
                ("Абдрахманова Алия Сериковна", None),
                (
                    "Совершенствование системы управления медицинской организацией",
                    None,
                ),
                (
                    "Магистерский проект на соискание степени магистра "
                    "здравоохранения",
                    None,
                ),
            ]
        )
        self.assertEqual(
            detect_dissertation_title_from_docx_bytes(blob),
            "Совершенствование системы управления медицинской организацией",
        )

    def test_govt_template_caps_title_above_prisuzhdenie_degree_marker_in_docx(self) -> None:
        blob = _build_docx_bytes(
            [
                ("НАО «Медицинский Университет Астана»", None),
                ("УДК: 303.446:618.1:338.465.4", None),
                ("МПК: G06Q30/01, G06Q30/012, G16H20/00", None),
                ("Амангелді Қуаныш Маратұлы", None),
                (
                    "СРАВНИТЕЛЬНЫЙ АНАЛИЗ ОКАЗАНИЯ ГИНЕКОЛОГИЧЕСКОЙ ПОМОЩИ "
                    "В КРУГЛОСУТОЧНОМ СТАЦИОНАРЕ НА ПЛАТНОЙ И "
                    "БЕЗВОЗМЕЗДНОЙ ОСНОВЕ В РАМКАХ КЛИНИКИ ТОО ЭКО",
                    None,
                ),
                (
                    "Диссертация на присуждение академической степени "
                    "магистра медицинских наук",
                    None,
                ),
            ]
        )
        self.assertEqual(
            detect_dissertation_title_from_docx_bytes(blob),
            "СРАВНИТЕЛЬНЫЙ АНАЛИЗ ОКАЗАНИЯ ГИНЕКОЛОГИЧЕСКОЙ ПОМОЩИ "
            "В КРУГЛОСУТОЧНОМ СТАЦИОНАРЕ НА ПЛАТНОЙ И "
            "БЕЗВОЗМЕЗДНОЙ ОСНОВЕ В РАМКАХ КЛИНИКИ ТОО ЭКО",
        )

    def test_govt_template_wrapped_mixed_case_kazakh_title_before_degree_line_in_docx(self) -> None:
        blob = _build_docx_bytes(
            [
                ("«Астана Медицина Университеті» КеАҚ", None),
                ("УДК: 614.2", None),
                ("МПК: A61B 5/12", None),
                ("Қожахметова Аружан Серікқызы", None),
                ("денсаулық сақтау ұйымындағы цифрлық", None),
                ("процестерді жетілдіру", None),
                ("магистрлік жоба (бейіндік бағыт үшін)", None),
            ]
        )
        self.assertEqual(
            detect_dissertation_title_from_docx_bytes(blob),
            "денсаулық сақтау ұйымындағы цифрлық процестерді жетілдіру",
        )

    def test_govt_template_accepts_degree_marker_with_schwa_in_docx(self) -> None:
        blob = _build_docx_bytes(
            [
                ("«Астана медициналық университеті» КеАҚ", None),
                ("УДК: 614.2:005:004.4", None),
                ("МПК: G16H10/00, G16H20/00", None),
                ("Акылбекова Айнур Оразбаевна", None),
                (
                    "Медициналық ұйымдағы икемді басқару əдістемелері: "
                    "операциялық тиімділікті арттыру үшін Agile жəне Scrum "
                    "əдісітерін еңгізуді зерттеу.",
                    None,
                ),
                ("7М10122 – «Денсаулық сақтау менеджменті бойынша МВА»", None),
                ("Магистр дəрежесін алу үшін ұсынылған жоба.", None),
            ]
        )
        self.assertEqual(
            detect_dissertation_title_from_docx_bytes(blob),
            "Медициналық ұйымдағы икемді басқару әдістемелері: "
            "операциялық тиімділікті арттыру үшін Agile және Scrum "
            "әдісітерін еңгізуді зерттеу",
        )

    def test_govt_template_accepts_mixed_latin_kazakh_letters_in_docx(self) -> None:
        blob = _build_docx_bytes(
            [
                ("«Астана медициналық университеті» КеАҚ", None),
                ("УДК: 614.2", None),
                ("МПК: A61B 5/12", None),
                ("Қожахметова Аружан Серікқызы", None),
                (
                    "Медициналық ұйымдаǵы üдерістерді жетıлдıру",
                    None,
                ),
                ("7М10116 - \"Қоғамдық денсаулық сақтау\"", None),
                ("Магистр дəрежесін алу үшін ұсынылған жоба.", None),
            ]
        )
        self.assertEqual(
            detect_dissertation_title_from_docx_bytes(blob),
            "Медициналық ұйымдағы үдерістерді жетілдіру",
        )

    def test_empty_blob(self) -> None:
        self.assertEqual(detect_dissertation_title_from_docx_bytes(b""), "")

    def test_corrupt_blob_returns_empty(self) -> None:
        self.assertEqual(
            detect_dissertation_title_from_docx_bytes(b"not a real docx"), ""
        )


# ---------------------------------------------------------------------------
# Язык диссертации — текст
# ---------------------------------------------------------------------------


_RUSSIAN_INTRO = (
    "Введение\n"
    "В настоящее время вопросы автоматизации обработки больших данных "
    "становятся всё более актуальными. Цель настоящей работы — исследовать "
    "методы машинного обучения и предложить новый подход к классификации "
    "документов на основе глубоких нейронных сетей. В работе используются "
    "современные алгоритмы и проводится сравнительный анализ с известными "
    "методами. Результаты исследования показывают, что предложенный подход "
    "позволяет достичь высокой точности на стандартных наборах данных."
)

_KAZAKH_INTRO = (
    "Кіріспе\n"
    "Қазіргі таңда үлкен деректерді өңдеу мәселелері өте өзекті болып "
    "табылады. Зерттеудің мақсаты — машиналық оқыту әдістерін зерттеу және "
    "құжаттарды жіктеудің жаңа тәсілін ұсыну. Жұмыста қазіргі алгоритмдер "
    "қолданылады және белгілі әдістермен салыстырмалы талдау жүргізіледі. "
    "Зерттеу нәтижелері ұсынылған тәсілдің стандартты деректер жиынтығында "
    "жоғары дәлдікке қол жеткізуге мүмкіндік беретінін көрсетеді."
)

_ENGLISH_INTRO = (
    "Introduction\n"
    "In recent years, the field of natural language processing has witnessed "
    "tremendous progress thanks to deep learning techniques. This thesis "
    "explores the application of transformer architectures to document "
    "classification tasks. We propose a novel approach that combines "
    "attention mechanisms with traditional feature engineering. Our "
    "experimental results demonstrate competitive performance on standard "
    "benchmarks while maintaining computational efficiency."
)


class DetectLanguageFromTextTests(unittest.TestCase):
    def test_normalize_kazakh_variant_letters(self) -> None:
        self.assertEqual(
            _normalize_kazakh_text("дəреже ǵылым üдеріс jetıстік óń"),
            "дәреже ғылым үдеріс jetістік өң",
        )

    def test_russian_text(self) -> None:
        self.assertEqual(detect_dissertation_language_from_text(_RUSSIAN_INTRO), LANGUAGE_RUSSIAN)

    def test_kazakh_text(self) -> None:
        self.assertEqual(detect_dissertation_language_from_text(_KAZAKH_INTRO), LANGUAGE_KAZAKH)

    def test_english_text(self) -> None:
        self.assertEqual(detect_dissertation_language_from_text(_ENGLISH_INTRO), LANGUAGE_ENGLISH)

    def test_empty_text(self) -> None:
        self.assertEqual(detect_dissertation_language_from_text(""), "")

    def test_text_without_letters(self) -> None:
        self.assertEqual(detect_dissertation_language_from_text("12345 67890 !!! ???"), "")

    def test_intro_marker_overrides_kazakh_abstract(self) -> None:
        # Казахская аннотация в начале + русское «Введение» — слайс должен
        # пойти после введения, поэтому язык = русский.
        text = (
            ("Аңдатпа\n" + ("Қазақша мәтін әңғқөұүһі " * 50)) + "\n"
            + "Введение\n" + (
                "Настоящая работа посвящена изучению вопросов обработки данных. "
                * 60
            )
        )
        self.assertEqual(detect_dissertation_language_from_text(text), LANGUAGE_RUSSIAN)

    def test_short_russian_text_uses_full_text(self) -> None:
        # < 500 символов — слайс возвращает весь текст, чтобы не потерять
        # короткие документы.
        short = "Это короткая русская диссертация без введения. Всего пара предложений."
        self.assertEqual(detect_dissertation_language_from_text(short), LANGUAGE_RUSSIAN)


class DetectLanguageFromGdocTests(unittest.TestCase):
    def test_russian_gdoc(self) -> None:
        document = _document([_paragraph(_RUSSIAN_INTRO + "\n")])
        self.assertEqual(detect_dissertation_language_from_gdoc(document), LANGUAGE_RUSSIAN)

    def test_kazakh_gdoc(self) -> None:
        document = _document([_paragraph(_KAZAKH_INTRO + "\n")])
        self.assertEqual(detect_dissertation_language_from_gdoc(document), LANGUAGE_KAZAKH)

    def test_empty_gdoc(self) -> None:
        self.assertEqual(detect_dissertation_language_from_gdoc({}), "")
        self.assertEqual(detect_dissertation_language_from_gdoc(_document([])), "")


class DetectLanguageFromDocxTests(unittest.TestCase):
    def test_russian_docx(self) -> None:
        blob = _build_docx_bytes([(_RUSSIAN_INTRO, None)])
        self.assertEqual(detect_dissertation_language_from_docx_bytes(blob), LANGUAGE_RUSSIAN)

    def test_kazakh_docx(self) -> None:
        blob = _build_docx_bytes([(_KAZAKH_INTRO, None)])
        self.assertEqual(detect_dissertation_language_from_docx_bytes(blob), LANGUAGE_KAZAKH)

    def test_empty_blob(self) -> None:
        self.assertEqual(detect_dissertation_language_from_docx_bytes(b""), "")


class WarnUnusualLanguageTests(unittest.TestCase):
    def test_warns_for_english(self) -> None:
        with self.assertLogs("magister_checking.dissertation_meta", level="WARNING") as cm:
            warn_if_unusual_language(LANGUAGE_ENGLISH, context="row=2")
        self.assertTrue(any("английский" in line.lower() or "english" in line.lower()
                            or LANGUAGE_ENGLISH in line for line in cm.output))

    def test_silent_for_russian(self) -> None:
        logger = logging.getLogger("magister_checking.dissertation_meta")
        with self.assertLogs(logger, level="WARNING") as cm:
            warn_if_unusual_language(LANGUAGE_RUSSIAN)
            # принудительно записать что-то, иначе assertLogs провалится при
            # отсутствии логов
            logger.warning("sentinel")
        self.assertEqual(len([line for line in cm.output if "sentinel" not in line]), 0)

    def test_silent_for_kazakh(self) -> None:
        logger = logging.getLogger("magister_checking.dissertation_meta")
        with self.assertLogs(logger, level="WARNING") as cm:
            warn_if_unusual_language(LANGUAGE_KAZAKH)
            logger.warning("sentinel")
        self.assertEqual(len([line for line in cm.output if "sentinel" not in line]), 0)

    def test_silent_for_empty(self) -> None:
        logger = logging.getLogger("magister_checking.dissertation_meta")
        with self.assertLogs(logger, level="WARNING") as cm:
            warn_if_unusual_language("")
            logger.warning("sentinel")
        self.assertEqual(len([line for line in cm.output if "sentinel" not in line]), 0)


if __name__ == "__main__":
    unittest.main()
